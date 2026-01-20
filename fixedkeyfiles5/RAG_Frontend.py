import warnings
warnings.filterwarnings("ignore", category=DeprecationWarning)
import sys
import os
import time
import csv
import json
import pandas as pd

try:
    import docx
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    docx = None

# 引入必要的 PyQt5 组件
from PyQt5.QtWidgets import (QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout, 
                             QLabel, QLineEdit, QPushButton, QTextEdit, QFileDialog, 
                             QMessageBox, QSplitter, QComboBox, QCheckBox, QRadioButton, 
                             QButtonGroup, QFrame, QGroupBox, QInputDialog, QDialog, QListWidget,
                             QSpinBox)

from PyQt5.QtCore import Qt, QSettings, QThread, pyqtSignal
from PyQt5.QtGui import QTextCursor, QColor

# 引入 Backend 逻辑
from RAG_Backend import RecallWorker
# 【Modified】从新文件引入本地模型加载器
from LocalModelLoader import LocalSemanticFilter

# ================= 默认航司列表 (三步走策略 - Step 1 数据源) =================
DEFAULT_AIRLINES = [
    "中国国际航空", "南方航空", "东方航空", "海南航空",
    "厦门航空", "四川航空", "深圳航空", "春秋航空",
    "吉祥航空", "首都航空", "山东航空", "天津航空",
    "上海航空", "祥鹏航空", "西部航空", "长龙航空",
    "Air China", "China Southern", "China Eastern"
]
AIRLINE_DICT_FILE = "airline_dict.txt"

# ================= 样式表 (Dark Mode) =================
STYLESHEET = """
QMainWindow { background-color: #2b2b2b; color: #e0e0e0; font-family: 'Segoe UI', 'Microsoft YaHei', sans-serif; }
QLabel { color: #aaaaaa; font-weight: bold; font-size: 13px; }
QLineEdit { background-color: #3c3c3c; color: #ffffff; border: 1px solid #555555; padding: 6px; border-radius: 4px; }
QTextEdit { background-color: #1e1e1e; color: #e0e0e0; border: 1px solid #444444; font-family: Consolas, monospace; font-size: 12px; }
QPushButton { background-color: #007acc; color: white; border: none; padding: 8px 16px; border-radius: 4px; font-weight: bold; font-size: 14px; }
QPushButton:hover { background-color: #005f9e; }
QPushButton:pressed { background-color: #004a80; }
QPushButton:disabled { background-color: #444444; color: #888888; }
/* Stop Button Style */
QPushButton#StopBtn { background-color: #d32f2f; }
QPushButton#StopBtn:hover { background-color: #b71c1c; }
QComboBox { background-color: #3c3c3c; color: white; border: 1px solid #555; padding: 5px; border-radius: 4px; }
QComboBox::drop-down { border: 0px; }
QSpinBox { background-color: #3c3c3c; color: white; border: 1px solid #555; padding: 5px; border-radius: 4px; }
QRadioButton { color: #e0e0e0; font-weight: bold; spacing: 5px; }
QRadioButton::indicator { width: 16px; height: 16px; }
QCheckBox { color: #e0e0e0; font-weight: bold; spacing: 5px; }
QGroupBox { border: 1px solid #555; margin-top: 10px; padding-top: 10px; font-weight: bold; color: #aaa; }
QGroupBox::title { subcontrol-origin: margin; subcontrol-position: top left; padding: 0 5px; }
QFrame#Divider { border: 1px solid #444444; }
/* 模式选择特定样式 */
QRadioButton#ModeSmart { color: #4facfe; }
QRadioButton#ModePrecise { color: #00f260; }
QRadioButton#ModeFuzzy { color: #ff9a9e; }
/* ListWidget for Dict Editor */
QListWidget { background-color: #333; color: white; border: 1px solid #555; }
"""

# ================= 航司字典编辑器窗口 =================
class AirlineDictEditor(QDialog):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowTitle("管理航司字典 (airline_dict.txt)")
        self.resize(400, 500)
        self.setStyleSheet(STYLESHEET)
        self.init_ui()
        self.load_data()

    def init_ui(self):
        layout = QVBoxLayout(self)
        
        # 输入区
        input_layout = QHBoxLayout()
        self.entry_new = QLineEdit()
        self.entry_new.setPlaceholderText("输入新航司名称...")
        btn_add = QPushButton("➕ 添加")
        btn_add.clicked.connect(self.add_item)
        btn_add.setFixedWidth(80)
        input_layout.addWidget(self.entry_new)
        input_layout.addWidget(btn_add)
        layout.addLayout(input_layout)

        # 列表区
        self.list_widget = QListWidget()
        layout.addWidget(self.list_widget)

        # 底部按钮
        btn_layout = QHBoxLayout()
        btn_del = QPushButton("❌ 删除选中")
        btn_del.setStyleSheet("background-color: #c0392b;")
        btn_del.clicked.connect(self.delete_item)
        
        btn_save = QPushButton("💾 保存并关闭")
        btn_save.setStyleSheet("background-color: #27ae60;")
        btn_save.clicked.connect(self.save_and_close)
        
        btn_layout.addWidget(btn_del)
        btn_layout.addStretch()
        btn_layout.addWidget(btn_save)
        layout.addLayout(btn_layout)

    def load_data(self):
        # 确保文件存在
        if not os.path.exists(AIRLINE_DICT_FILE):
            try:
                with open(AIRLINE_DICT_FILE, "w", encoding="utf-8") as f:
                    for airline in DEFAULT_AIRLINES:
                        f.write(airline + "\n")
            except Exception:
                pass
        
        # 读取
        try:
            with open(AIRLINE_DICT_FILE, "r", encoding="utf-8") as f:
                lines = f.readlines()
                for line in lines:
                    text = line.strip()
                    if text:
                        self.list_widget.addItem(text)
        except Exception as e:
            QMessageBox.warning(self, "Error", f"读取字典失败: {str(e)}")

    def add_item(self):
        text = self.entry_new.text().strip()
        if not text:
            return
        # 查重
        items = [self.list_widget.item(i).text() for i in range(self.list_widget.count())]
        if text in items:
            QMessageBox.warning(self, "提示", "该关键词已存在")
            return
        
        self.list_widget.addItem(text)
        self.entry_new.clear()

    def delete_item(self):
        row = self.list_widget.currentRow()
        if row >= 0:
            self.list_widget.takeItem(row)

    def save_and_close(self):
        items = [self.list_widget.item(i).text() for i in range(self.list_widget.count())]
        try:
            with open(AIRLINE_DICT_FILE, "w", encoding="utf-8") as f:
                for item in items:
                    f.write(item + "\n")
            QMessageBox.information(self, "成功", "航司字典已更新！")
            self.accept()
        except Exception as e:
            QMessageBox.critical(self, "错误", f"保存失败: {str(e)}")

# ================= 后台模型加载线程 =================
class ModelInitWorker(QThread):
    finished_signal = pyqtSignal(object) # 返回 LocalSemanticFilter 实例

    def run(self):
        print(">>> 正在后台初始化本地小模型 (LocalSemanticFilter)...")
        filter_instance = LocalSemanticFilter()
        self.finished_signal.emit(filter_instance)

# ================= 主界面 =================
class RAGRecallApp(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("RAG 工业级全流程 (Recall + RRF Fusion + Multi-Model Support)")
        self.resize(1400, 950) 
        self.setStyleSheet(STYLESHEET)
        
        self.settings = QSettings("MyCorp", "RAGRecall_Final_v11_LocalModel")
        self.cached_results = []
        self.cached_summary = ""
        self.cached_query = ""
        self.worker = None 
        
        # 本地小模型实例
        self.local_filter_instance = None
        
        # 初始化字典文件
        self.ensure_airline_dict()
        
        self.init_ui()
        
        # 启动后台加载模型
        self.model_loader = ModelInitWorker()
        self.model_loader.finished_signal.connect(self.on_model_loaded)
        self.model_loader.start()
        
    def on_model_loaded(self, instance):
        self.local_filter_instance = instance
        if instance.is_loaded:
            self.log("✅ 本地小模型加载成功！已准备好进行语义预筛选。")
        else:
            self.log("⚠️ 本地小模型加载失败或库缺失，相关功能将不可用。")

    def ensure_airline_dict(self):
        """确保航司字典文件存在"""
        if not os.path.exists(AIRLINE_DICT_FILE):
            try:
                with open(AIRLINE_DICT_FILE, "w", encoding="utf-8") as f:
                    for airline in DEFAULT_AIRLINES:
                        f.write(airline + "\n")
            except:
                pass

    def init_ui(self):
        central_widget = QWidget()
        self.setCentralWidget(central_widget)
        
        main_layout = QHBoxLayout(central_widget)
        main_layout.setSpacing(10)
        
        # === Left Widget ===
        left_widget = QWidget()
        left_layout = QVBoxLayout(left_widget)
        
        # 1. DB & JSON Inputs
        db_layout = QHBoxLayout()
        self.db_path_edit = QLineEdit()
        self.db_path_edit.setText(self.settings.value("last_db_path", ""))
        btn_db = QPushButton("📂 向量库")
        btn_db.clicked.connect(self.browse_db)
        db_layout.addWidget(QLabel("Vector DB:"))
        db_layout.addWidget(self.db_path_edit)
        db_layout.addWidget(btn_db)
        left_layout.addLayout(db_layout)
        
        json_layout = QHBoxLayout()
        self.json_path_edit = QLineEdit()
        self.json_path_edit.setText(self.settings.value("last_json_path", ""))
        btn_json = QPushButton("📄 PageIndex")
        btn_json.clicked.connect(self.browse_json)
        json_layout.addWidget(QLabel("Structure JSON:"))
        json_layout.addWidget(self.json_path_edit)
        json_layout.addWidget(btn_json)
        left_layout.addLayout(json_layout)
        
        # 2. Search Mode & Model Selection
        mode_layout = QHBoxLayout()
        
        # 2.1 策略选择
        mode_layout.addWidget(QLabel("🔍 策略:"))
        self.mode_group = QButtonGroup(self)
        
        self.radio_smart = QRadioButton("🔵 智能")
        self.radio_smart.setObjectName("ModeSmart")
        self.radio_smart.setChecked(True)
        self.mode_group.addButton(self.radio_smart, 1)
        mode_layout.addWidget(self.radio_smart)
        
        self.radio_precise = QRadioButton("🟢 精准")
        self.radio_precise.setObjectName("ModePrecise")
        self.mode_group.addButton(self.radio_precise, 2)
        mode_layout.addWidget(self.radio_precise)
        
        self.radio_fuzzy = QRadioButton("🟡 模糊")
        self.radio_fuzzy.setObjectName("ModeFuzzy")
        self.mode_group.addButton(self.radio_fuzzy, 3)
        mode_layout.addWidget(self.radio_fuzzy)
        
        mode_layout.addSpacing(15)

        # 2.2 Doc Type Selection & Dynamic Throttling
        mode_layout.addWidget(QLabel("📂 文档偏好:"))
        self.combo_doc_type = QComboBox()
        self.combo_doc_type.addItems([
            "不指定类型",
            "数据表",
            "公司公文",
            "书籍/教材",   
            "长篇论文",    
            "技术文档",
            "法律条文",
            "LLM OCR文档",
            "LLM生成总结"
        ])
        self.combo_doc_type.setFixedWidth(120)
        mode_layout.addWidget(self.combo_doc_type)

        # 2.3 Model Selection
        mode_layout.addWidget(QLabel("🧠 模型:"))
        self.combo_model = QComboBox()
        self.combo_model.addItems([
            "DeepSeek-R1", 
            "DeepSeek-V3", 
            "X1-70B-thinking", 
            "X1-70B-fast"
        ])
        self.combo_model.setFixedWidth(140)
        mode_layout.addWidget(self.combo_model)
        
        mode_layout.addStretch()
        left_layout.addLayout(mode_layout)

        # 3. Knowledge Config & Advanced Controls
        config_group = QGroupBox("🔧 知识库与高级检索配置")
        config_layout = QVBoxLayout(config_group)
        
        # 3.1 Stopwords & Buttons
        sw_btn_layout = QHBoxLayout()
        
        self.stopwords_edit = QTextEdit()
        self.stopwords_edit.setPlaceholderText("在此输入停用词，以逗号分隔...")
        self.stopwords_edit.setMaximumHeight(35)
        sw_btn_layout.addWidget(self.stopwords_edit, 2)

        self.sw_name_edit = QLineEdit()
        self.sw_name_edit.setPlaceholderText("Config Name")
        self.sw_name_edit.setFixedWidth(80)
        sw_btn_layout.addWidget(self.sw_name_edit)
        
        btn_sw_save = QPushButton("💾 Save")
        btn_sw_save.clicked.connect(self.save_stopwords)
        btn_sw_save.setFixedHeight(28)
        sw_btn_layout.addWidget(btn_sw_save)
        
        btn_sw_load = QPushButton("📂 Load")
        btn_sw_load.clicked.connect(self.import_stopwords)
        btn_sw_load.setFixedHeight(28)
        sw_btn_layout.addWidget(btn_sw_load)
        
        # 航司字典按钮
        btn_airline = QPushButton("✈️ 航司字典")
        btn_airline.clicked.connect(self.open_airline_editor)
        btn_airline.setFixedHeight(28)
        btn_airline.setStyleSheet("background-color: #8e44ad; color: white;")
        sw_btn_layout.addWidget(btn_airline)
        
        config_layout.addLayout(sw_btn_layout)

        # 3.2 Advanced Search Control (New Requirements)
        adv_layout = QHBoxLayout()
        
        # FAISS 勾选
        self.chk_use_faiss = QCheckBox("⚡ 启用 FAISS 加速")
        self.chk_use_faiss.setChecked(True)
        self.chk_use_faiss.setStyleSheet("color: #00e676;") 
        adv_layout.addWidget(self.chk_use_faiss)
        
        adv_layout.addSpacing(15)

        # 小模型勾选 (Requirement 1)
        self.chk_local_model = QCheckBox("🧠 启用本地小模型过滤")
        self.chk_local_model.setChecked(True) # 默认勾选
        self.chk_local_model.setStyleSheet("color: #ffa502;")
        self.chk_local_model.setToolTip("勾选后，会在发送给API Rerank前使用本地小模型进行语义预筛选，节省成本并提高同义词召回率。")
        adv_layout.addWidget(self.chk_local_model)
        
        adv_layout.addStretch()
        
        # 切片数控制 (Requirement 2)
        adv_layout.addWidget(QLabel("向量召回数:"))
        self.spin_vector_k = QSpinBox()
        self.spin_vector_k.setRange(5, 20)
        self.spin_vector_k.setValue(10)
        self.spin_vector_k.setFixedWidth(50)
        adv_layout.addWidget(self.spin_vector_k)

        adv_layout.addWidget(QLabel("硬查询数:"))
        self.spin_json_k = QSpinBox()
        self.spin_json_k.setRange(5, 20)
        self.spin_json_k.setValue(10)
        self.spin_json_k.setFixedWidth(50)
        adv_layout.addWidget(self.spin_json_k)

        config_layout.addLayout(adv_layout)
        
        left_layout.addWidget(config_group)

        # 4. Query Input
        left_layout.addWidget(QLabel("用户查询 (Query):"))
        self.query_input = QTextEdit()
        self.query_input.setPlaceholderText("请输入问题 (如：查找航班 JMU，或询问某操作流程)...")
        self.query_input.setMaximumHeight(60)
        left_layout.addWidget(self.query_input)
        
        # 5. Search & Stop Buttons
        btn_layout = QHBoxLayout()
        
        self.btn_search = QPushButton("🚀 执行全流程 (Recall -> RRF -> Summary)")
        self.btn_search.setFixedHeight(45)
        self.btn_search.setStyleSheet("background-color: #2da44e; font-size: 15px;")
        self.btn_search.clicked.connect(self.start_recall)
        btn_layout.addWidget(self.btn_search)
        
        self.btn_stop = QPushButton("🛑 停止运行")
        self.btn_stop.setObjectName("StopBtn")
        self.btn_stop.setFixedHeight(45)
        self.btn_stop.setEnabled(False)
        self.btn_stop.clicked.connect(self.stop_recall)
        btn_layout.addWidget(self.btn_stop)
        
        left_layout.addLayout(btn_layout)
        
        # 6. Result Displays
        left_layout.addWidget(QLabel("🤖 智能总结 (Thinking + Answer):"))
        self.summary_display = QTextEdit()
        self.summary_display.setReadOnly(True)
        self.summary_display.setStyleSheet("""
            QTextEdit {
                background-color: #252526; 
                color: #dcdcaa; 
                font-family: 'Segoe UI', sans-serif; 
                font-size: 14px; 
                border: 1px solid #007acc;
                line-height: 1.6;
            }
        """)
        self.summary_display.setMinimumHeight(250)
        left_layout.addWidget(self.summary_display)

        left_layout.addWidget(QLabel("📚 RRF Fused Context (Top-12 Unique):"))
        self.result_display = QTextEdit()
        self.result_display.setReadOnly(True)
        self.result_display.setStyleSheet("font-family: Consolas; font-size: 12px; color: #aaddff;")
        left_layout.addWidget(self.result_display)
        
        # 7. Export Area
        export_layout = QHBoxLayout()
        export_layout.addWidget(QLabel("导出格式:"))
        
        self.combo_format = QComboBox()
        self.combo_format.addItems(["xlsx", "csv", "txt", "docx", "md"])
        self.combo_format.setFixedWidth(100)
        export_layout.addWidget(self.combo_format)
        
        self.btn_export = QPushButton("💾 导出结果")
        self.btn_export.setStyleSheet("background-color: #d2691e;")
        self.btn_export.clicked.connect(self.export_data)
        export_layout.addWidget(self.btn_export)
        
        export_layout.addStretch() 
        left_layout.addLayout(export_layout)

        # === Right Console ===
        right_widget = QWidget()
        right_layout = QVBoxLayout(right_widget)
        right_layout.addWidget(QLabel("📟 System Console"))
        self.console_output = QTextEdit()
        self.console_output.setReadOnly(True)
        self.console_output.setStyleSheet("background-color: #111; color: #0f0; font-family: Consolas;")
        right_layout.addWidget(self.console_output)
        
        splitter = QSplitter(Qt.Horizontal)
        splitter.addWidget(left_widget)
        splitter.addWidget(right_widget)
        splitter.setSizes([900, 400])
        main_layout.addWidget(splitter)
        
        # 加载上次的 Stopwords
        last_sw = self.settings.value("last_stopwords", "")
        if last_sw:
            self.stopwords_edit.setText(last_sw)

    # ================= UI 交互逻辑 =================

    def browse_db(self):
        path, _ = QFileDialog.getOpenFileName(self, "选择数据库", "", "SQLite DB (*.db);;All Files (*.*)")
        if path:
            self.db_path_edit.setText(path)
            self.settings.setValue("last_db_path", path)

    def browse_json(self):
        path, _ = QFileDialog.getOpenFileName(self, "选择 JSON", "", "JSON Files (*.json);;All Files (*.*)")
        if path:
            self.json_path_edit.setText(path)
            self.settings.setValue("last_json_path", path)

    def log(self, msg):
        self.console_output.append(msg)
        cursor = self.console_output.textCursor()
        cursor.movePosition(QTextCursor.End)
        self.console_output.setTextCursor(cursor)

    def update_summary(self, text):
        """流式更新总结文本"""
        self.cached_summary = text 
        self.summary_display.setMarkdown(text) 
        cursor = self.summary_display.textCursor()
        cursor.movePosition(QTextCursor.End)
        self.summary_display.setTextCursor(cursor)

    def display_results(self, results):
        self.cached_results = results 
        html = ""
        for item in results:
            score_text = f"{item['final_score']:.4f}"
            debug_info = item.get('debug_score', '')
            local_score = item.get('local_score', None)
            
            source = item.get('source', 'VECTOR')
            
            if "JSON" in source:
                source_style = "background-color: #00f260; color: #111; padding: 2px 5px; border-radius: 3px; font-weight: bold;"
            else:
                source_style = "background-color: #007acc; color: white; padding: 2px 5px; border-radius: 3px;"
            
            # 显示额外信息
            extra_badges = ""
            if local_score is not None:
                extra_badges += f" <span style='background-color:#e67e22;color:white;padding:1px 4px;border-radius:3px;'>LocalSim: {local_score:.3f}</span>"
            
            score_color = "#00ff00" 
            
            html += f"""
            <div style='border-bottom: 1px solid #555; padding: 12px; margin-bottom: 8px;'>
                <span style='color: #888; font-weight:bold;'>Rank #{item['rank']}</span> | 
                <span style='{source_style}'>{source}</span> | 
                <span style='color: {score_color}; font-weight: bold;'>RRF Score: {score_text}</span> 
                {extra_badges}
                <span style='color: #aaa; font-size:11px;'>[{debug_info}]</span><br>
                <div style='margin-top:5px; color: #ffcc00;'><b>[Section Path]</b> {item['path']}</div>
                <div style='margin-top:5px; background-color: #222; padding: 8px; border-left: 3px solid #2da44e; white-space: pre-wrap;'>
{item['content'][:200]}...
                </div>
            </div>
            """
        self.result_display.setHtml(html)

    # ================= 知识库配置功能 (Stopwords + Airline) =================
    
    def open_airline_editor(self):
        """打开航司字典编辑器"""
        editor = AirlineDictEditor(self)
        editor.exec_() # 模态运行

    def get_current_stopwords(self):
        text = self.stopwords_edit.toPlainText().strip()
        if not text:
            return []
        return [w.strip() for w in text.replace('，', ',').split(',') if w.strip()]
    
    def get_airline_list(self):
        """读取航司字典"""
        airlines = []
        if os.path.exists(AIRLINE_DICT_FILE):
            try:
                with open(AIRLINE_DICT_FILE, 'r', encoding='utf-8') as f:
                    airlines = [line.strip() for line in f if line.strip()]
            except:
                pass
        return airlines

    def save_stopwords(self):
        name = self.sw_name_edit.text().strip()
        if not name:
            QMessageBox.warning(self, "Warning", "Please enter a profile name.")
            return
        
        sw_list = self.get_current_stopwords()
        filename = f"stopwords_{name}.json"
        
        try:
            with open(filename, 'w', encoding='utf-8') as f:
                json.dump({"name": name, "stopwords": sw_list}, f, ensure_ascii=False, indent=2)
            self.log(f"💾 Stopwords profile '{name}' saved to {filename}")
            QMessageBox.information(self, "Success", f"Saved to {filename}")
        except Exception as e:
            self.log(f"❌ Save failed: {str(e)}")

    def import_stopwords(self):
        path, _ = QFileDialog.getOpenFileName(self, "Import Stopwords JSON", "", "JSON Files (*.json);;All Files (*.*)")
        if path:
            try:
                with open(path, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                
                if isinstance(data, list):
                    sw_list = data
                elif isinstance(data, dict) and "stopwords" in data:
                    sw_list = data["stopwords"]
                else:
                    raise ValueError("Invalid JSON format")
                
                self.stopwords_edit.setText(", ".join(sw_list))
                self.log(f"📂 Imported {len(sw_list)} words from {os.path.basename(path)}")
            except Exception as e:
                self.log(f"❌ Import failed: {str(e)}")

    # ================= 核心流程控制 =================

    def start_recall(self):
        # 1. 基础校验
        db_path = self.db_path_edit.text().strip()
        json_path = self.json_path_edit.text().strip()
        query = self.query_input.toPlainText().strip()
        
        if not db_path or not os.path.exists(db_path):
            QMessageBox.warning(self, "Error", "无效的数据库路径")
            return
        if not query:
            return

        # 2. 获取参数
        try:
            mode_id = self.mode_group.checkedId()
            search_mode = "smart" 
            if mode_id == 2: search_mode = "precise"
            elif mode_id == 3: search_mode = "fuzzy"

            summary_model = self.combo_model.currentText()
            doc_type = self.combo_doc_type.currentText()
            stopwords = self.get_current_stopwords()
            airline_names = self.get_airline_list()
            
            # 获取高级选项
            use_faiss = self.chk_use_faiss.isChecked()
            enable_local_filter = self.chk_local_model.isChecked()
            
            # 获取切片数量
            limit_vector = self.spin_vector_k.value()
            limit_json = self.spin_json_k.value()
            
            self.cached_query = query
            self.cached_results = []
            self.cached_summary = ""
            
            # UI 状态更新
            self.btn_search.setEnabled(False)
            self.btn_stop.setEnabled(True) 
            self.result_display.clear()
            self.summary_display.clear() 
            self.console_output.clear()
            
            mode_text = {"smart": "🔵 智能融合", "precise": "🟢 精准查表", "fuzzy": "🟡 模糊咨询"}[search_mode]
            self.log(f"🚀 初始化任务... | 策略: {mode_text} | 模型: {summary_model}")
            self.log(f"ℹ️ 加速模式: {'FAISS' if use_faiss else 'Brute-Force JSON'}")
            
            if enable_local_filter:
                self.log(f"🧠 本地小模型过滤: 已启用 (Vector Top {limit_vector})")
            else:
                self.log(f"⚪ 本地小模型过滤: 未启用 (Vector Top {limit_vector})")
                
            self.log(f"ℹ️ Hard Query Limit: {limit_json}")
            
            # 实例化后端 Worker
            self.worker = RecallWorker(
                query_text=query, 
                db_path=db_path, 
                json_path=json_path, 
                search_mode=search_mode, 
                summary_model=summary_model, 
                doc_type=doc_type, 
                stopwords=stopwords, 
                airline_names=airline_names,
                use_faiss=use_faiss,
                enable_local_filter=enable_local_filter,
                local_filter_instance=self.local_filter_instance,
                limit_vector=limit_vector,
                limit_json=limit_json
            )
            
            self.worker.log_signal.connect(self.log)
            self.worker.result_signal.connect(self.display_results)
            self.worker.summary_signal.connect(self.update_summary) 
            self.worker.finish_signal.connect(self.on_finished)
            self.worker.start()

        except Exception as e:
            # 兜底：防止前端逻辑崩溃
            QMessageBox.critical(self, "System Error", f"启动失败: {str(e)}")
            self.btn_search.setEnabled(True)
            self.btn_stop.setEnabled(False)

    def stop_recall(self):
        if self.worker and self.worker.isRunning():
            self.log("🛑 用户点击停止，正在请求 Backend 中断...")
            self.worker.stop()
            self.btn_stop.setEnabled(False) 
            self.btn_stop.setText("Stopping...")

    def on_finished(self, success):
        self.btn_search.setEnabled(True)
        self.btn_stop.setEnabled(False)
        self.btn_stop.setText("🛑 停止运行")
        
        if success:
            self.log("✅ 全流程结束")
        else:
            self.log("❌ 流程被中断或发生错误")
            if not self.cached_summary:
                self.summary_display.setText("[ Process Stopped / Failed ]")

    # ================= 导出功能 =================
    def export_data(self):
        if not self.cached_summary and not self.cached_results:
            QMessageBox.warning(self, "提示", "当前没有可导出的结果，请先执行查询。")
            return

        fmt = self.combo_format.currentText()
        timestamp = time.strftime("%Y%m%d_%H%M%S")
        filename = f"export_{timestamp}.{fmt}"
        save_path = os.path.join(os.getcwd(), filename)

        try:
            self.log(f"💾 正在导出为 {fmt} ...")
            
            if fmt == "xlsx":
                if pd is None:
                    raise ImportError("缺少 pandas 或 openpyxl 库")
                
                data_rows = []
                for item in self.cached_results:
                    data_rows.append({
                        "Rank": item['rank'],
                        "Source": item.get('source', 'VECTOR'),
                        "RRF Score": item['final_score'],
                        "Local Score": item.get('local_score', ''),
                        "Section Path": item['path'],
                        "Content": item['content']
                    })
                
                df = pd.DataFrame(data_rows)
                
                with pd.ExcelWriter(save_path, engine='openpyxl') as writer:
                    summary_df = pd.DataFrame([["DeepSeek R1 Summary"], [self.cached_summary], [""]])
                    summary_df.to_excel(writer, sheet_name='Report', index=False, header=False, startrow=0)
                    
                    pd.DataFrame([["Top Results"]]).to_excel(writer, sheet_name='Report', index=False, header=False, startrow=4)
                    df.to_excel(writer, sheet_name='Report', index=False, startrow=6)
            
            elif fmt == "csv":
                with open(save_path, 'w', newline='', encoding='utf-8-sig') as f:
                    writer = csv.writer(f)
                    writer.writerow(["=== DeepSeek R1 Summary ==="])
                    writer.writerow([self.cached_summary])
                    writer.writerow([])
                    writer.writerow(["=== Top Results ==="])
                    writer.writerow(["Rank", "Source", "RRF Score", "Local Score", "Section Path", "Content"])
                    for item in self.cached_results:
                        writer.writerow([
                            item['rank'],
                            item.get('source', 'VECTOR'),
                            f"{item['final_score']:.4f}",
                            item.get('local_score', ''),
                            item['path'],
                            item['content']
                        ])

            elif fmt == "txt":
                with open(save_path, 'w', encoding='utf-8') as f:
                    f.write(f"Query: {self.cached_query}\n")
                    f.write("="*50 + "\n")
                    f.write("DeepSeek R1 Summary:\n")
                    f.write("="*50 + "\n")
                    clean_summary = self.cached_summary.replace("**", "").replace(">", "")
                    f.write(clean_summary + "\n\n")
                    f.write("="*50 + "\n")
                    f.write("Top Results:\n")
                    f.write("="*50 + "\n")
                    for item in self.cached_results:
                        f.write(f"[Rank #{item['rank']}] [{item.get('source','VECTOR')}] RRF: {item['final_score']:.4f}\n")
                        if 'local_score' in item:
                            f.write(f"Local Score: {item['local_score']}\n")
                        f.write(f"Path: {item['path']}\n")
                        f.write(f"Content:\n{item['content']}\n")
                        f.write("-" * 30 + "\n")

            elif fmt == "md":
                with open(save_path, 'w', encoding='utf-8') as f:
                    f.write(f"# RAG Query Report\n\n")
                    f.write(f"**Query:** {self.cached_query}\n\n")
                    f.write(f"## 🤖 DeepSeek R1 Summary\n\n")
                    f.write(self.cached_summary + "\n\n")
                    f.write(f"## 📚 Top Results\n\n")
                    for item in self.cached_results:
                        f.write(f"### Rank #{item['rank']} [{item.get('source','VECTOR')}] (RRF: {item['final_score']:.4f})\n")
                        f.write(f"**Path:** `{item['path']}`\n\n")
                        f.write(f"**Content:**\n\n")
                        content_block = item['content'].replace('\n', '\n> ')
                        f.write(f"> {content_block}\n\n")
                        f.write("---\n")

            elif fmt == "docx":
                if docx is None:
                    raise ImportError("缺少 python-docx 库")
                
                doc = docx.Document()
                doc.add_heading('RAG Analysis Report', 0)
                
                p = doc.add_paragraph()
                p.add_run('Query: ').bold = True
                p.add_run(self.cached_query)
                
                doc.add_heading('DeepSeek R1 Summary', level=1)
                doc.add_paragraph(self.cached_summary)
                
                doc.add_heading('Top Results', level=1)
                
                for item in self.cached_results:
                    p_header = doc.add_paragraph()
                    run = p_header.add_run(f"Rank #{item['rank']} | [{item.get('source','VECTOR')}] | RRF: {item['final_score']:.4f}")
                    run.bold = True
                    run.font.color.rgb = docx.shared.RGBColor(0, 100, 0)
                    
                    if 'local_score' in item and item['local_score'] is not None:
                         p_header.add_run(f" | Local: {item['local_score']:.3f}")

                    p_path = doc.add_paragraph()
                    p_path.add_run("Path: ").bold = True
                    p_path.add_run(item['path']).italic = True
                    
                    p_content = doc.add_paragraph(item['content'])
                    p_content.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    doc.add_paragraph("-" * 40)

            self.log(f"✅ 导出成功: {save_path}")
            QMessageBox.information(self, "成功", f"文件已导出至:\n{save_path}")

        except ImportError as ie:
            self.log(f"❌ 导出失败 (库缺失): {str(ie)}")
            QMessageBox.critical(self, "错误", f"导出失败，缺少必要库:\n{str(ie)}")
        except Exception as e:
            self.log(f"❌ 导出异常: {str(e)}")
            import traceback
            self.log(traceback.format_exc())
            QMessageBox.critical(self, "错误", f"导出过程中发生错误:\n{str(e)}")

if __name__ == "__main__":
    app = QApplication(sys.argv)
    window = RAGRecallApp()
    window.show()
    sys.exit(app.exec_())