import sys
import json
import os
import subprocess
import csv

from PyQt5.QtWidgets import (QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout, 
                             QLabel, QLineEdit, QPushButton, QTextEdit, QComboBox, 
                             QFileDialog, QMessageBox, QFrame, QTabWidget, QSplitter, 
                             QListWidget, QListWidgetItem, QShortcut)
from PyQt5.QtCore import Qt, QThread, pyqtSignal
from PyQt5.QtGui import QColor, QPalette, QFont, QTextCursor, QKeySequence, QTextCharFormat

# --- 尝试导入外部依赖 ---
# 1. 尝试导入 AI 可视化窗口 (来自 pgui.py)
try:
    from ai_visual_window import AIVisualWindow
    HAS_VISUAL_WINDOW = True
except ImportError:
    HAS_VISUAL_WINDOW = False
    # 如果缺失，创建一个哑类防止报错
    class AIVisualWindow(QWidget):
        def add_stream_char(self, char): pass

# 2. 尝试导入 docx (来自 pgirecallwindow.py)
try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

# 3. 尝试导入 pandas (来自 pgirecallwindow.py)
try:
    import pandas as pd
    HAS_PANDAS = True
except ImportError:
    HAS_PANDAS = False

CONFIG_FILE = "gui_configs.json"

# === 全局统一样式表 (合并了两个文件的风格) ===
GLOBAL_STYLESHEET = """
QMainWindow, QWidget {
    background-color: #0d1117;
    color: #c9d1d9;
}
QTabWidget::pane {
    border: 1px solid #30363d;
    background: #0d1117;
}
QTabBar::tab {
    background: #161b22;
    color: #8b949e;
    padding: 10px 20px;
    border: 1px solid #30363d;
    border-bottom: none;
    border-top-left-radius: 4px;
    border-top-right-radius: 4px;
    margin-right: 2px;
}
QTabBar::tab:selected {
    background: #0d1117;
    color: #58a6ff;
    border-bottom: 2px solid #58a6ff; /* 高亮当前 Tab */
}
QTabBar::tab:hover {
    background: #21262d;
}

/* 标签与标题 */
QLabel {
    color: #58a6ff; /* 统一使用蓝色系标题，或保留pgui的青色 */
    font-family: 'Segoe UI', 'Microsoft YaHei', sans-serif;
    font-weight: bold;
}

/* 输入框 */
QLineEdit {
    background-color: #161b22;
    border: 1px solid #30363d;
    border-radius: 4px;
    color: #c9d1d9;
    padding: 6px;
    font-family: 'Consolas', 'Microsoft YaHei';
}
QLineEdit:focus {
    border: 1px solid #58a6ff;
    background-color: #0d1117;
}

/* 按钮 */
QPushButton {
    background-color: #238636;
    color: white;
    border: none;
    padding: 6px 12px;
    border-radius: 6px;
    font-weight: bold;
    font-size: 13px;
}
QPushButton:hover {
    background-color: #2ea043;
}
QPushButton:pressed {
    background-color: #1a6329;
}

/* 特殊按钮: Visual Switch */
QPushButton#VisualBtn {
    background-color: #1f6feb;
    border: 1px solid #1f6feb;
}
QPushButton#VisualBtn:hover {
    background-color: #388bfd;
}

/* 文本编辑区 */
QTextEdit {
    background-color: #0d1117;
    border: 1px solid #30363d;
    color: #c9d1d9; 
    font-family: 'Consolas', 'Microsoft YaHei', monospace;
    font-size: 13px;
    line-height: 1.5;
}

/* 下拉框 */
QComboBox {
    background-color: #161b22;
    color: #c9d1d9;
    border: 1px solid #30363d;
    padding: 5px;
    border-radius: 4px;
}
QComboBox::drop-down {
    border: none;
}

/* 列表控件 (Recall Window) */
QListWidget { 
    background-color: #0d1117; 
    border: 1px solid #30363d; 
    border-radius: 6px;
    color: #c9d1d9; 
    font-size: 14px; 
    padding: 5px;
}
QListWidget::item { padding: 5px; }
QListWidget::item:selected { background-color: #1f6feb; border-radius: 4px; color: white; }

/* 分隔条 */
QSplitter::handle { background-color: #30363d; }
"""

# =================================================================================
# 模块 1: 索引构建 (原 pgui.py 的逻辑)
# =================================================================================

class WorkerThread(QThread):
    log_signal = pyqtSignal(str)      
    stream_signal = pyqtSignal(str)   

    def __init__(self, command):
        super().__init__()
        self.command = command
        self.line_buffer = ""

    def run(self):
        process = subprocess.Popen(
            self.command,
            stdout=subprocess.PIPE,
            stderr=subprocess.STDOUT, 
            shell=True,
            text=True,
            encoding='utf-8',
            errors='replace',
            bufsize=0 
        )

        while True:
            char = process.stdout.read(1)
            if not char and process.poll() is not None:
                break
            if char:
                self.process_char(char)
        
        self.flush_buffer()
        process.wait()

    def flush_buffer(self):
        if self.line_buffer:
            line = self.line_buffer.strip()
            if line:
                self.emit_log_line(line)
            self.line_buffer = ""

    def process_char(self, char):
        self.line_buffer += char
        if char == "\n":
            line = self.line_buffer.strip()
            if line: 
                if line.startswith("DEBUG_AI_CHAR:"):
                    try:
                        content = line.split("DEBUG_AI_CHAR:", 1)[1]
                        self.stream_signal.emit(content)
                    except: pass
                else:
                    self.emit_log_line(line)
            self.line_buffer = ""

    def emit_log_line(self, line):
        if "[SUCCESS]" in line:
            formatted_line = f"<span style='color:#00FF00; font-weight:bold; font-size:13px;'>{line}</span>"
        elif "[ERROR]" in line or "Exception" in line or "Traceback" in line or "Error" in line:
            formatted_line = f"<span style='color:#FF3333; font-weight:bold;'>{line}</span>"
        elif "[INFO]" in line:
            formatted_line = f"<span style='color:#33CCFF;'>{line}</span>"
        elif "[Warning]" in line:
            formatted_line = f"<span style='color:#FFFF00;'>{line}</span>"
        else:
            formatted_line = line
        self.log_signal.emit(formatted_line)

class IndexerTab(QWidget): # 从 QMainWindow 改为 QWidget
    def __init__(self):
        super().__init__()
        # 初始化可视化窗口
        if HAS_VISUAL_WINDOW:
            self.visual_window = AIVisualWindow()
        else:
            self.visual_window = QWidget() # 占位
        
        self.configs = self.load_configs()
        self.init_ui()

    def init_ui(self):
        layout = QVBoxLayout(self)
        layout.setSpacing(15)
        layout.setContentsMargins(20, 20, 20, 20)

        # === Header Area ===
        header_layout = QHBoxLayout()
        title_label = QLabel("INDEXER & PROCESSOR")
        title_label.setStyleSheet("font-size: 20px; color: #00ffcc; letter-spacing: 2px;")
        header_layout.addWidget(title_label)
        header_layout.addStretch()
        layout.addLayout(header_layout)

        # === Config Section ===
        cfg_frame = QFrame()
        cfg_frame.setStyleSheet("background-color: #161b22; border-radius: 8px; padding: 10px;")
        cfg_layout = QHBoxLayout(cfg_frame)
        
        self.cb_configs = QComboBox()
        self.cb_configs.addItems(self.configs.keys())
        self.cb_configs.currentTextChanged.connect(self.load_selected_config)
        
        btn_save = QPushButton("💾 SAVE CONFIG")
        btn_save.clicked.connect(self.save_config)
        btn_save.setStyleSheet("background-color: #21262d; border: 1px solid #30363d;")

        cfg_layout.addWidget(QLabel("CONFIGURATION:"))
        cfg_layout.addWidget(self.cb_configs, 1)
        cfg_layout.addWidget(btn_save)
        layout.addWidget(cfg_frame)

        # === Input Section ===
        input_layout = QVBoxLayout()
        
        # PDF Selection
        file_layout = QHBoxLayout()
        self.edit_pdf = QLineEdit()
        self.edit_pdf.setPlaceholderText("Select PDF document path...")
        btn_file = QPushButton("📂 BROWSE")
        btn_file.clicked.connect(self.get_file)
        file_layout.addWidget(QLabel("DOCUMENT:"))
        file_layout.addWidget(self.edit_pdf, 1)
        file_layout.addWidget(btn_file)
        input_layout.addLayout(file_layout)
        
        # Model Selection
        model_layout = QHBoxLayout()
        self.edit_model = QLineEdit("DeepSeek-V3")
        model_layout.addWidget(QLabel("AI MODEL:"))
        model_layout.addWidget(self.edit_model, 1)
        input_layout.addLayout(model_layout)
        
        layout.addLayout(input_layout)

        # === Action Buttons ===
        btn_layout = QHBoxLayout()
        
        self.btn_run = QPushButton("🚀 INITIALIZE INDEXING")
        self.btn_run.setFixedHeight(45)
        self.btn_run.clicked.connect(self.start_task)
        
        self.btn_visual = QPushButton("👁️ VISUALIZER: OFF")
        self.btn_visual.setObjectName("VisualBtn")
        self.btn_visual.setCheckable(True)
        self.btn_visual.setFixedHeight(45)
        self.btn_visual.clicked.connect(self.toggle_visual_window)
        if not HAS_VISUAL_WINDOW:
             self.btn_visual.setEnabled(False)
             self.btn_visual.setText("👁️ VISUALIZER (Missing)")

        btn_layout.addWidget(self.btn_run, 2)
        btn_layout.addWidget(self.btn_visual, 1)
        layout.addLayout(btn_layout)

        # === Console Output ===
        layout.addWidget(QLabel("SYSTEM LOGS:"))
        self.txt_console = QTextEdit()
        self.txt_console.setReadOnly(True)
        # 这里特别指定一下控制台的字体颜色，保持Matrix风格
        self.txt_console.setStyleSheet("color: #00ff99; font-family: 'Consolas', monospace; font-size: 12px;")
        layout.addWidget(self.txt_console)

    def toggle_visual_window(self):
        if not HAS_VISUAL_WINDOW: return
        if self.btn_visual.isChecked():
            self.visual_window.show()
            self.btn_visual.setText("👁️ VISUALIZER: ON")
            # 尝试移动到主窗口右侧，但这里 self 是 tab，需要获取 window
            window = self.window()
            if window:
                geo = window.geometry()
                self.visual_window.move(geo.x() + geo.width() + 10, geo.y())
        else:
            self.visual_window.hide()
            self.btn_visual.setText("👁️ VISUALIZER: OFF")

    def load_configs(self):
        if os.path.exists(CONFIG_FILE):
            try:
                with open(CONFIG_FILE, 'r', encoding='utf-8') as f: return json.load(f)
            except: pass
        return {"Default": {"pdf": "", "model": "DeepSeek-V3", "pages": "3"}}

    def save_config(self):
        name = self.cb_configs.currentText() or "NewConfig"
        self.configs[name] = {"pdf": self.edit_pdf.text(), "model": self.edit_model.text(), "pages": "3"}
        with open(CONFIG_FILE, 'w', encoding='utf-8') as f: json.dump(self.configs, f)
        QMessageBox.information(self, "System", "Configuration Saved Successfully.")

    def load_selected_config(self, name):
        if name in self.configs:
            c = self.configs[name]
            self.edit_pdf.setText(c.get('pdf',''))
            self.edit_model.setText(c.get('model','DeepSeek-V3'))

    def get_file(self):
        f, _ = QFileDialog.getOpenFileName(self, "Select PDF", "", "*.pdf")
        if f: self.edit_pdf.setText(f)

    def append_log(self, text):
        self.txt_console.append(text)
        cursor = self.txt_console.textCursor()
        cursor.movePosition(QTextCursor.End)
        self.txt_console.setTextCursor(cursor)

    def start_task(self):
        pdf_path = self.edit_pdf.text()
        if not pdf_path:
            QMessageBox.warning(self, "Error", "Please select a PDF file first.")
            return

        py_exe = sys.executable
        # 注意：这里假设 run_pageindex.py 在同一目录下
        cmd = f'"{py_exe}" -u run_pageindex.py --pdf_path "{pdf_path}" --model "{self.edit_model.text()}" --toc-check-pages 3'
        
        self.txt_console.clear()
        self.txt_console.append(f"<span style='color:#FFFF00'>[SYSTEM] Initializing subprocess...</span>")
        
        self.worker = WorkerThread(cmd)
        self.worker.log_signal.connect(self.append_log)
        
        if HAS_VISUAL_WINDOW:
            self.worker.stream_signal.connect(self.visual_window.add_stream_char)
            if not self.btn_visual.isChecked():
                self.btn_visual.click()
            
        self.worker.start()


# =================================================================================
# 模块 2: 知识召回 (原 pgirecallwindow.py 的逻辑)
# =================================================================================

class RecallTab(QWidget): # 从 QMainWindow 改为 QWidget
    def __init__(self):
        super().__init__()
        self.data = None
        self.all_nodes = [] 
        self.last_loaded_path = None
        
        self.init_ui()
        self.setup_shortcuts()

    def init_ui(self):
        layout = QVBoxLayout(self)

        # --- Top Bar ---
        top_bar = QHBoxLayout()
        
        self.btn_load = QPushButton("📂 加载索引")
        self.btn_load.clicked.connect(self.open_file_dialog)
        
        self.btn_refresh = QPushButton("🔄 刷新")
        self.btn_refresh.setToolTip("重新加载当前文件并显示全部内容")
        self.btn_refresh.clicked.connect(self.refresh_data)
        
        self.edit_search = QLineEdit()
        self.edit_search.setPlaceholderText("🔍 输入关键词进行全局内容召回...")
        self.edit_search.returnPressed.connect(self.search_content)
        
        self.btn_search = QPushButton("执行召回")
        self.btn_search.clicked.connect(self.search_content)

        self.combo_export = QComboBox()
        self.combo_export.addItems(["DOCX (Word)", "TXT (文本)", "CSV (表格)", "XLSX (Excel)"])
        self.combo_export.setFixedWidth(120)
        
        self.btn_export = QPushButton("💾 导出全文")
        self.btn_export.clicked.connect(self.export_data)
        
        top_bar.addWidget(self.btn_load)
        top_bar.addWidget(self.btn_refresh)
        top_bar.addWidget(self.edit_search, 4)
        top_bar.addWidget(self.btn_search)
        top_bar.addSpacing(20)
        top_bar.addWidget(QLabel("格式:"))
        top_bar.addWidget(self.combo_export)
        top_bar.addWidget(self.btn_export)
        
        layout.addLayout(top_bar)

        # --- Splitter ---
        splitter = QSplitter(Qt.Horizontal)
        
        # Left: Results
        left_widget = QWidget()
        left_layout = QVBoxLayout(left_widget)
        left_layout.setContentsMargins(0, 0, 0, 0)
        left_layout.addWidget(QLabel("召回结果列表:"))
        self.list_results = QListWidget()
        self.list_results.itemClicked.connect(self.display_node_detail)
        left_layout.addWidget(self.list_results)
        
        # Right: Details
        right_widget = QWidget()
        right_layout = QVBoxLayout(right_widget)
        right_layout.setContentsMargins(0, 0, 0, 0)
        right_layout.setSpacing(5)
        
        right_layout.addWidget(QLabel("详情预览:"))
        self.txt_header = QTextEdit()
        self.txt_header.setReadOnly(True)
        self.txt_header.setMaximumHeight(110)
        self.txt_header.setStyleSheet("border: none; background-color: #0d1117;") 
        right_layout.addWidget(self.txt_header)

        # Detail Search
        search_bar_layout = QHBoxLayout()
        search_label = QLabel("🔎 正文查找:")
        search_label.setStyleSheet("color: #8b949e; font-size: 12px;")
        
        self.edit_inner_search = QLineEdit()
        self.edit_inner_search.setPlaceholderText("在此处输入文本，按回车高亮显示 (Ctrl+F)")
        self.edit_inner_search.textChanged.connect(self.highlight_text_in_detail)
        self.edit_inner_search.setStyleSheet("""
            background-color: #21262d; border: 1px solid #30363d; 
            color: #ffd700; font-weight: bold;
        """)
        
        search_bar_layout.addWidget(search_label)
        search_bar_layout.addWidget(self.edit_inner_search)
        right_layout.addLayout(search_bar_layout)

        self.txt_content = QTextEdit()
        self.txt_content.setReadOnly(True)
        right_layout.addWidget(self.txt_content)
        
        splitter.addWidget(left_widget)
        splitter.addWidget(right_widget)
        splitter.setStretchFactor(0, 1)
        splitter.setStretchFactor(1, 3)
        
        layout.addWidget(splitter)

    def setup_shortcuts(self):
        # 注意: Shortcut 需要绑定到组件，否则在Tab切换时可能冲突，这里绑定到self
        self.shortcut_find = QShortcut(QKeySequence("Ctrl+F"), self)
        self.shortcut_find.activated.connect(self.focus_inner_search)

    def focus_inner_search(self):
        # 确保只有当前Tab显示时才生效
        if self.isVisible():
            self.edit_inner_search.setFocus()
            self.edit_inner_search.selectAll()

    def open_file_dialog(self):
        file_path, _ = QFileDialog.getOpenFileName(self, "选择索引文件", "", "JSON Files (*.json);;All Files (*)")
        if file_path:
            self.load_file_content(file_path)

    def refresh_data(self):
        if self.last_loaded_path and os.path.exists(self.last_loaded_path):
            self.edit_search.clear()
            self.load_file_content(self.last_loaded_path)
            self.txt_content.append(f"\n🔄 已刷新数据，显示全部内容。")
        else:
            QMessageBox.warning(self, "无法刷新", "尚未加载文件或文件路径已失效。")

    def load_file_content(self, file_path):
        try:
            # === 关键修复：使用 utf-8-sig 自动处理 BOM ===
            with open(file_path, 'r', encoding='utf-8-sig') as f:
                self.data = json.load(f)
            
            self.all_nodes = []

            # === JSON 结构解析逻辑（已兼容新旧格式） ===
            if isinstance(self.data, dict):
                if 'structure' in self.data:
                    structure = self.data['structure']
                else:
                    structure = [self.data]
            elif isinstance(self.data, list):
                structure = self.data
            else:
                raise ValueError("JSON 文件格式不被识别")

            self._flatten_structure(structure)
            # ======================================

            self.last_loaded_path = file_path
            
            self.txt_content.setText(f"✅ 已成功加载文件: {os.path.basename(file_path)}\n📊 共解析出 {len(self.all_nodes)} 个知识节点。\n\n请在上方搜索框输入关键词进行召回。")
            self.txt_header.clear()
            self.edit_inner_search.clear()

            self.list_results.clear()
            for node in self.all_nodes:
                self._add_item_to_list(node)
                
        except Exception as e:
            import traceback
            self.txt_content.setText(f"❌ 加载失败: {str(e)}\n\n{traceback.format_exc()}")
            QMessageBox.critical(self, "错误", f"读取文件失败:\n{str(e)}")

    def _flatten_structure(self, structure):
        if not structure:
            return
        for item in structure:
            self.all_nodes.append(item)
            if 'nodes' in item and isinstance(item['nodes'], list):
                self._flatten_structure(item['nodes'])

    def search_content(self):
        query = self.edit_search.text().strip().lower()
        self.list_results.clear()
        
        if not query:
            for node in self.all_nodes:
                self._add_item_to_list(node)
            return
            
        results_found = 0
        for node in self.all_nodes:
            title = node.get('title', '').lower()
            text = node.get('text', '').lower()
            if query in title or query in text:
                self._add_item_to_list(node)
                results_found += 1
        
        if results_found > 0:
            self.txt_content.setText(f"🔍 查询关键字: '{query}'\n✅ 成功召回到 {results_found} 个匹配章节。")
        else:
            self.txt_content.setText(f"⚠️ 未找到包含 '{query}' 的内容。")

    def _add_item_to_list(self, node):
        title = node.get('title', '无标题节点')
        display_title = (title[:40] + '...') if len(title) > 40 else title
        item = QListWidgetItem(display_title)
        item.setToolTip(title)
        item.setData(Qt.UserRole, node)
        self.list_results.addItem(item)

    def display_node_detail(self, item):
        node = item.data(Qt.UserRole)
        if node:
            start = node.get('start_index', '-')
            end = node.get('end_index', '-')
            
            header_html = f"""
            <h2 style='color: #58a6ff; margin-bottom: 5px;'>{node.get('title', '未命名章节')}</h2>
            <div style='background-color: #21262d; padding: 5px; border-radius: 5px;'>
                <span style='color: #8b949e; font-weight: bold;'>📄 物理页码:</span> 
                <span style='color: #c9d1d9;'>第 {start} - {end} 页</span>
                &nbsp;&nbsp;&nbsp;|&nbsp;&nbsp;&nbsp;
                <span style='color: #8b949e; font-weight: bold;'>🆔 Node ID:</span> 
                <span style='color: #c9d1d9;'>{node.get('node_id', 'N/A')}</span>
            </div>
            """
            self.txt_header.setHtml(header_html)
            
            raw_text = node.get('text', '')
            if not raw_text:
                raw_text = "（该节点无正文内容）"
            
            self.txt_content.setPlainText(raw_text)
            
            if self.edit_inner_search.text():
                self.highlight_text_in_detail()

    def highlight_text_in_detail(self):
        search_str = self.edit_inner_search.text()
        cursor = self.txt_content.textCursor()
        cursor.select(QTextCursor.Document)
        format = QTextCharFormat()
        format.setBackground(Qt.transparent)
        cursor.setCharFormat(format)
        
        if not search_str:
            return

        highlight_format = QTextCharFormat()
        highlight_format.setBackground(QColor("#d29922"))
        highlight_format.setForeground(QColor("black"))

        cursor = self.txt_content.textCursor()
        cursor.setPosition(0)
        
        while True:
            cursor = self.txt_content.document().find(search_str, cursor)
            if cursor.isNull():
                break
            cursor.mergeCharFormat(highlight_format)

    def export_data(self):
        if not self.all_nodes:
            QMessageBox.warning(self, "无数据", "当前没有加载任何数据，无法导出。")
            return

        file_format = self.combo_export.currentText()
        default_ext = ""
        filter_str = ""
        
        if "DOCX" in file_format:
            default_ext = ".docx"
            filter_str = "Word Document (*.docx)"
        elif "TXT" in file_format:
            default_ext = ".txt"
            filter_str = "Text File (*.txt)"
        elif "CSV" in file_format:
            default_ext = ".csv"
            filter_str = "CSV File (*.csv)"
        elif "XLSX" in file_format:
            default_ext = ".xlsx"
            filter_str = "Excel File (*.xlsx)"

        save_path, _ = QFileDialog.getSaveFileName(self, "导出文件", f"export_data{default_ext}", filter_str)
        if not save_path:
            return

        try:
            if "DOCX" in file_format:
                self._export_docx(save_path)
            elif "TXT" in file_format:
                self._export_txt(save_path)
            elif "CSV" in file_format:
                self._export_tabular(save_path, is_csv=True)
            elif "XLSX" in file_format:
                self._export_tabular(save_path, is_csv=False)
            
            QMessageBox.information(self, "成功", f"文件已成功导出至:\n{save_path}")
            
        except Exception as e:
            QMessageBox.critical(self, "导出失败", f"发生错误: {str(e)}\n如果是库缺失，请运行: pip install python-docx pandas openpyxl")

    def _export_docx(self, path):
        if not HAS_DOCX:
            raise ImportError("未安装 python-docx 库。")
        doc = Document()
        doc.add_heading('索引数据导出', 0)
        for node in self.all_nodes:
            title = node.get('title', '无标题')
            text = node.get('text', '')
            page_info = f"页码: {node.get('start_index', '-')} - {node.get('end_index', '-')}"

            doc.add_heading(title, level=1)
            p = doc.add_paragraph()
            run = p.add_run(page_info)
            run.italic = True
            doc.add_paragraph(text)
            doc.add_paragraph("-" * 20)
        doc.save(path)

    def _export_txt(self, path):
        with open(path, 'w', encoding='utf-8') as f:
            for node in self.all_nodes:
                f.write(f"【标题】: {node.get('title', '无标题')}\n")
                f.write(f"【页码】: {node.get('start_index', '-')} - {node.get('end_index', '-')}\n")
                f.write(f"【正文】:\n{node.get('text', '')}\n")
                f.write("-" * 50 + "\n\n")

    def _export_tabular(self, path, is_csv=True):
        if not HAS_PANDAS:
            raise ImportError("未安装 pandas 库。")
        data_list = []
        for node in self.all_nodes:
            data_list.append({
                "Node ID": node.get('node_id', ''),
                "Title": node.get('title', ''),
                "Start Page": node.get('start_index', ''),
                "End Page": node.get('end_index', ''),
                "Content": node.get('text', '')
            })
        df = pd.DataFrame(data_list)
        if is_csv:
            df.to_csv(path, index=False, encoding='utf-8-sig')
        else:
            df.to_excel(path, index=False)

# =================================================================================
# 核心主窗口: UnifiedMainWindow
# =================================================================================

class UnifiedMainWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("PageIndex Pro - Integrated Suite (DeepSeek适配版)")
        self.resize(1200, 900)
        
        # 应用统一样式
        self.setStyleSheet(GLOBAL_STYLESHEET)
        
        # 初始化 Tab Widget
        self.tabs = QTabWidget()
        self.setCentralWidget(self.tabs)
        
        # 添加页面
        self.tab_indexer = IndexerTab()
        self.tab_recall = RecallTab()
        
        self.tabs.addTab(self.tab_indexer, "🔧 索引构建 (Indexer)")
        self.tabs.addTab(self.tab_recall, "🔎 知识召回 (Recall)")
        
        # 设置 Tab 字体大小
        tab_bar = self.tabs.tabBar()
        font = tab_bar.font()
        font.setPointSize(11)
        font.setBold(True)
        tab_bar.setFont(font)

if __name__ == "__main__":
    app = QApplication(sys.argv)
    app.setStyle('Fusion') # 融合风格，对CSS支持较好
    window = UnifiedMainWindow()
    window.show()
    sys.exit(app.exec_())