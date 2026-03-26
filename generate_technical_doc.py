import os
import datetime
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def set_font_style(run, name, size, bold=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color

def create_tech_doc():
    doc = Document()

    # 设置默认中文字体
    doc.styles['Normal'].font.name = 'SimSun'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')

    # Title
    title = doc.add_heading('', 0)
    run = title.add_run('DeepSeek-VectifyAI-PageIndex 技术全案文档')
    set_font_style(run, 'Microsoft YaHei', 24, True)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f'生成日期: {datetime.date.today()}')
    doc.add_paragraph('版本: 1.5')
    doc.add_paragraph('作者: AI 自动生成文档助手')

    # 1. 代码结构 (Code Structure)
    doc.add_heading('1. 代码结构 (Code Structure)', level=1)
    doc.add_paragraph('项目的整体组织结构严谨，采用了典型的模块化设计：')
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '目录/文件'
    hdr_cells[1].text = '功能描述'

    struct_data = [
        ('pageindex/', '核心业务逻辑包，包含文档解析、树结构生成、工具类等'),
        ('  - page_index.py', '主逻辑：实现基于 LLM 的 PDF 目录检测、物理索引匹配和递归节点细化'),
        ('  - utils.py', '底层支撑：封装了 DeepSeek/Qwen API 调用、Token 计数、JSON 提取、PDF 文本处理'),
        ('  - config.yaml', '项目静态配置，定义了节点大小限制、API 默认参数等'),
        ('pgui.py', '基于 PyQt5 的全功能桌面 GUI，集成了索引生成、RAG 向量化准备等功能'),
        ('run_pageindex.py', '生产环境命令行入口，支持完整的 TOC 检测和验证逻辑'),
        ('run_pageindex_simple.py', '快速启动脚本，跳过耗时的 TOC 检测，适用于快速验证'),
        ('RAG_FASTAPI_WEB_toolkit/', 'Web 服务端组件，提供基于 FastAPI 的 RAG 接口'),
        ('docs/', '系统化的文档体系（PRD, SRS, 部署手册等）'),
        ('results/', '存放生成的结构化索引 JSON 文件，命名通常为 {pdf_name}_full.json'),
        ('tests/', '包含示例 PDF 文件和对应的测试结果')
    ]
    for path, desc in struct_data:
        row_cells = table.add_row().cells
        row_cells[0].text = path
        row_cells[1].text = desc

    # 2. RAG切片核心逻辑
    doc.add_heading('2. RAG切片核心逻辑 (Core RAG Slicing Logic)', level=1)
    doc.add_paragraph('PageIndex 采用了创新的“Agentic Tree Indexing”策略，而非传统的“固定块切片”：')
    doc.add_paragraph(
        '• 语义分级切片：不按字符数强行切割，而是通过 LLM 识别文档的“自然段落”和“标题层级”。\n'
        '• 递归细化逻辑：如果一个章节内容过多（超过 20,000 Tokens 或 10 页），系统会递归地将其作为“子文档”进行二次解析。\n'
        '• 物理页码对齐：每个切片都严格绑定了原始 PDF 的物理页码（physical_index），确保检索结果可回溯、可验证。\n'
        '• 摘要驱动：为每个层级节点生成语义摘要，检索时先通过摘要定位子树，再深入具体内容。'
    )

    # 3. 功能地图 (Feature Map)
    doc.add_heading('3. 功能地图 (Feature Map)', level=1)
    features = [
        'PageIndex 生成器：将 PDF/Markdown 转换为层级 JSON。',
        'RAG 向量化增强：为解析后的节点生成“判断型摘要”，优化召回率。',
        '多模式召回引擎：支持智能路由（Smart）、精确匹配（Precise）和模糊语义（Fuzzy）。',
        '可视化监控：GUI 界面实时展示 AI 推理过程中的“思考”内容。',
        '多模型支持：原生适配 DeepSeek-V3, Qwen-Plus, GPT-4o 等模型。'
    ]
    for feat in features:
        doc.add_paragraph(feat, style='List Bullet')

    # 4. 关键算法 (Key Algorithms)
    doc.add_heading('4. 关键算法 (Key Algorithms)', level=1)
    doc.add_paragraph('• TOC 自动化识别算法：结合 LLM 的思维链（CoT）判断页面是否具备目录特征，提取标题与页码的映射关系。')
    doc.add_paragraph('• 物理索引修正算法（Offset Calculation）：通过抽样验证标题在实际页面中的位置，计算逻辑页码与物理页码的差值（Offset）。')
    doc.add_paragraph('• 递归树构建算法：深度优先搜索（DFS）遍历文档，动态调整 start_index 和 end_index，构建完美的平衡树结构。')

    # 5. 数据结构 (Data Structure)
    doc.add_heading('5. 数据结构 (Data Structure)', level=1)
    doc.add_paragraph('系统核心索引对象（Node）的定义：')
    code_para = doc.add_paragraph()
    run = code_para.add_run(
        '{\n'
        '  "title": "String, 章节标题",\n'
        '  "node_id": "String, 4位填充ID",\n'
        '  "start_index": "Int, 起始物理页码",\n'
        '  "end_index": "Int, 结束物理页码",\n'
        '  "summary": "String, LLM 生成的章节综述",\n'
        '  "nodes": "Array[Node], 子层级节点",\n'
        '  "text": "String, (可选) 原始文本内容"\n'
        '}'
    )
    set_font_style(run, 'Consolas', 10)

    # 6. 函数调用图 (Function Call Graph)
    doc.add_heading('6. 函数调用图 (Function Call Graph)', level=1)
    doc.add_paragraph(
        '用户触发 -> pgui.py (GUI 事件)\n'
        '  -> run_pageindex.py (参数解析)\n'
        '    -> page_index.py: page_index_main()\n'
        '      -> tree_parser()\n'
        '        -> check_toc() -> find_toc_pages()\n'
        '        -> meta_processor()\n'
        '          -> process_toc_with_page_numbers() \n'
        '          -> verify_toc() -> check_title_appearance()\n'
        '        -> process_large_node_recursively() [递归调用]\n'
        '      -> add_node_text()\n'
        '      -> generate_summaries_for_structure()'
    )

    # 7. 安全性分析、可扩展性还有性能
    doc.add_heading('7. 安全性分析、可扩展性还有性能', level=1)
    doc.add_paragraph('安全性：采用局部解析策略，不上传原始二进制 PDF，仅处理提取后的文本。API 访问受环境变量保护。')
    doc.add_paragraph('性能：引入 AsyncIO 并发机制，在摘要生成阶段实现 10x 以上的加速。Slim Version 机制防止了内存溢出。')
    doc.add_paragraph('可扩展性：通过适配器模式，可轻松接入不同的 LLM 供应商和 Vector DB。')

    # 8. 总结和建议
    doc.add_heading('8. 总结和建议', level=1)
    doc.add_paragraph('总结：PageIndex 为解决长文档 RAG 丢失上下文的问题提供了极佳的范式。')
    doc.add_paragraph('建议：1. 增强对多列布局 PDF 的解析鲁棒性；2. 集成本地嵌入模型（如 BGE-M3）以降低 API 开销。')

    # 9. 产品 MD 文档
    doc.add_heading('9. 产品 MD 文档 (PRD Summary)', level=1)
    doc.add_paragraph('产品定位：面向专业领域（金融、审计、技术）的高精度文档 AI 分析平台。')

    # 10 & 13. User Story
    doc.add_heading('10. User Stories (Set 1)', level=1)
    doc.add_paragraph('1. 作为一个法务人员，我需要从 500 页的合同中快速找到违约条款，PageIndex 帮我直接定位到了第 412 页。')
    doc.add_heading('11. User Stories (Set 2)', level=1)
    doc.add_paragraph('2. 作为一个开发者，我需要将公司内部手册接入 RAG，传统切片导致答案断章取义，PageIndex 保持了语义完整。')

    # 12. SRS 文档
    doc.add_heading('12. SRS 文档 (Software Requirements Specification)', level=1)
    doc.add_paragraph('需求：支持 1GB 以上 PDF 处理；检索准确率 95%+ (基于 FinanceBench 测试)；支持多角色权限控制。')

    # 14. 开发文档
    doc.add_heading('14. 开发文档 (Development Guide)', level=1)
    doc.add_paragraph('基础环境：Python 3.9+。核心依赖：PyQt5, PyPDF2, Tiktoken。开发模式：模块化开发，建议使用 venv 环境。')

    # 15. 本地小模型应用部署技术文档
    doc.add_heading('15. 本地小模型应用部署技术文档', level=1)
    doc.add_paragraph('部署建议：使用 Ollama 运行 DeepSeek-R1-Distill-Qwen-7B，设置 API_BASE=http://localhost:11434/v1。系统可实现全离线运行。')

    # 16, 17, 18. Maps
    doc.add_heading('16. Mind Map / Mermaid / Markmap', level=1)
    doc.add_paragraph('以下为逻辑图代码（可粘贴至 Mermaid Live Editor 查看）：')
    run = doc.add_paragraph().add_run(
        'mindmap\n'
        '  root((PageIndex))\n'
        '    Parsing\n'
        '      PDF Extraction\n'
        '      Markdown Parsing\n'
        '    Indexing\n'
        '      TOC Detection\n'
        '      Recursive Slicing\n'
        '      Semantic Summaries\n'
        '    RAG Integration\n'
        '      Vectorization\n'
        '      Smart Recall\n'
    )
    set_font_style(run, 'Consolas', 9)

    # 19. Requirements.txt
    doc.add_heading('17. Requirements.txt 内容清单', level=1)
    doc.add_paragraph('PyQt5, openai, tiktoken, requests, urllib3, jieba, numpy, python-dotenv, PyYAML, PyPDF2, python-docx, langchain, chromadb, faiss-cpu')

    # 20. 软件使用手册
    doc.add_heading('18. 软件使用手册 (User Manual)', level=1)
    doc.add_paragraph('第一步：配置 .env 密钥。\n第二步：运行 python pgui.py。\n第三步：上传 PDF 并点击解析。\n第四步：在 RAG 界面进行交互问答。')

    # Save
    save_path = 'DeepSeek-VectifyAI-PageIndex_Technical_Full_Doc.docx'
    doc.save(save_path)
    print(f'Document saved to {save_path}')

if __name__ == '__main__':
    create_tech_doc()
