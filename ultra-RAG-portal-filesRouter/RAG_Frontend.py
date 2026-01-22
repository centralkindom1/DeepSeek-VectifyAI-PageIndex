import warnings
warnings.filterwarnings("ignore", category=DeprecationWarning)
import sys
import os
import time
import csv
import json
import sqlite3  # Added for Routing Database
import pandas as pd
import gc # 引入垃圾回收，用于模型切换时的内存清理
import requests # Added for API
import urllib3 # Added for SSL warning suppression
import numpy as np # Added for Vector Calculation

# 抑制 SSL 警告
urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)

# ================= 路由 API 配置 =================
ROUTING_API_KEY = "sk-fXM4W0CdcKnNp3NVDfF85f2b90284b11AfDdF9F5627f627b"
ROUTING_EMBED_URL = "https://aiplus.airchina.com.cn:18080/v1/embeddings"
ROUTING_MODEL_NAME = "bge-m3"

# 尝试导入 psutil 用于系统监控
try:
    import psutil
    HAS_PSUTIL = True
except ImportError:
    HAS_PSUTIL = False

try:
    import docx
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    docx = None

# 引入必要的 PyQt5 组件
from PyQt5.QtWidgets import (QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout, 
                             QLabel, QLineEdit, QPushButton, QTextEdit, QFileDialog, 
                             QMessageBox, QSplitter, QComboBox, QCheckBox, QRadioButton, 
                             QButtonGroup, QFrame, QGroupBox, QInputDialog, QDialog, QListWidget,
                             QSpinBox, QProgressBar)

from PyQt5.QtCore import Qt, QSettings, QThread, pyqtSignal
from PyQt5.QtGui import QTextCursor, QColor

# 引入 Backend 逻辑
from RAG_Backend import RecallWorker

# 【Modified】引入两个本地模型加载器
from LocalModelLoader import LocalSemanticFilter
from LocalBGELoader import BGESemanticFilter

# ================= 默认航司列表 =================
DEFAULT_AIRLINES = [
    "中国国际航空", "南方航空", "东方航空", "海南航空",
    "厦门航空", "四川航空", "深圳航空", "春秋航空",
    "吉祥航空", "首都航空", "山东航空", "天津航空",
    "上海航空", "祥鹏航空", "西部航空", "长龙航空",
    "Air China", "China Southern", "China Eastern"
]
AIRLINE_DICT_FILE = "airline_dict.txt"

# ================= 样式表 (Dark Mode) =================
STYLESHEET = """
QMainWindow { background-color: #2b2b2b; color: #e0e0e0; font-family: 'Segoe UI', 'Microsoft YaHei', sans-serif; }
QLabel { color: #aaaaaa; font-weight: bold; font-size: 13px; }
QLineEdit { background-color: #3c3c3c; color: #ffffff; border: 1px solid #555555; padding: 6px; border-radius: 4px; }
QTextEdit { background-color: #1e1e1e; color: #e0e0e0; border: 1px solid #444444; font-family: Consolas, monospace; font-size: 12px; }
QPushButton { background-color: #007acc; color: white; border: none; padding: 8px 16px; border-radius: 4px; font-weight: bold; font-size: 14px; }
QPushButton:hover { background-color: #005f9e; }
QPushButton:pressed { background-color: #004a80; }
QPushButton:disabled { background-color: #444444; color: #888888; }
/* Stop Button Style */
QPushButton#StopBtn { background-color: #d32f2f; }
QPushButton#StopBtn:hover { background-color: #b71c1c; }
QComboBox { background-color: #3c3c3c; color: white; border: 1px solid #555; padding: 5px; border-radius: 4px; }
QComboBox::drop-down { border: 0px; }
QSpinBox { background-color: #3c3c3c; color: white; border: 1px solid #555; padding: 5px; border-radius: 4px; }
QRadioButton { color: #e0e0e0; font-weight: bold; spacing: 5px; }
QRadioButton::indicator { width: 16px; height: 16px; }
QCheckBox { color: #e0e0e0; font-weight: bold; spacing: 5px; }
QGroupBox { border: 1px solid #555; margin-top: 10px; padding-top: 10px; font-weight: bold; color: #aaa; }
QGroupBox::title { subcontrol-origin: margin; subcontrol-position: top left; padding: 0 5px; }
QFrame#Divider { border: 1px solid #444444; }
/* 进度条样式 */
QProgressBar { border: 1px solid #555; border-radius: 3px; text-align: center; color: white; background-color: #333; }
QProgressBar::chunk { background-color: #2da44e; width: 10px; margin: 0.5px; }
QListWidget { background-color: #333; color: white; border: 1px solid #555; }
"""

# ================= 航司字典编辑器窗口 =================
class AirlineDictEditor(QDialog):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowTitle("管理航司字典 (airline_dict.txt)")
        self.resize(400, 500)
        self.setStyleSheet(STYLESHEET)
        self.init_ui()
        self.load_data()

    def init_ui(self):
        layout = QVBoxLayout(self)
        
        # 输入区
        input_layout = QHBoxLayout()
        self.entry_new = QLineEdit()
        self.entry_new.setPlaceholderText("输入新航司名称...")
        btn_add = QPushButton("➕ 添加")
        btn_add.clicked.connect(self.add_item)
        btn_add.setFixedWidth(80)
        input_layout.addWidget(self.entry_new)
        input_layout.addWidget(btn_add)
        layout.addLayout(input_layout)

        # 列表区
        self.list_widget = QListWidget()
        layout.addWidget(self.list_widget)

        # 底部按钮
        btn_layout = QHBoxLayout()
        btn_del = QPushButton("❌ 删除选中")
        btn_del.setStyleSheet("background-color: #c0392b;")
        btn_del.clicked.connect(self.delete_item)
        
        btn_save = QPushButton("💾 保存并关闭")
        btn_save.setStyleSheet("background-color: #27ae60;")
        btn_save.clicked.connect(self.save_and_close)
        
        btn_layout.addWidget(btn_del)
        btn_layout.addStretch()
        btn_layout.addWidget(btn_save)
        layout.addLayout(btn_layout)

    def load_data(self):
        if not os.path.exists(AIRLINE_DICT_FILE):
            try:
                with open(AIRLINE_DICT_FILE, "w", encoding="utf-8") as f:
                    for airline in DEFAULT_AIRLINES:
                        f.write(airline + "\n")
            except Exception:
                pass
        
        try:
            with open(AIRLINE_DICT_FILE, "r", encoding="utf-8") as f:
                lines = f.readlines()
                for line in lines:
                    text = line.strip()
                    if text:
                        self.list_widget.addItem(text)
        except Exception as e:
            QMessageBox.warning(self, "Error", f"读取字典失败: {str(e)}")

    def add_item(self):
        text = self.entry_new.text().strip()
        if not text: return
        items = [self.list_widget.item(i).text() for i in range(self.list_widget.count())]
        if text in items:
            QMessageBox.warning(self, "提示", "该关键词已存在")
            return
        self.list_widget.addItem(text)
        self.entry_new.clear()

    def delete_item(self):
        row = self.list_widget.currentRow()
        if row >= 0:
            self.list_widget.takeItem(row)

    def save_and_close(self):
        items = [self.list_widget.item(i).text() for i in range(self.list_widget.count())]
        try:
            with open(AIRLINE_DICT_FILE, "w", encoding="utf-8") as f:
                for item in items:
                    f.write(item + "\n")
            QMessageBox.information(self, "成功", "航司字典已更新！")
            self.accept()
        except Exception as e:
            QMessageBox.critical(self, "错误", f"保存失败: {str(e)}")

# ================= 系统监视线程 (Win7 保护) =================
class SystemMonitorWorker(QThread):
    cpu_signal = pyqtSignal(float)
    
    def run(self):
        while True:
            try:
                if HAS_PSUTIL:
                    # 采集 CPU 占用率，阻塞 2 秒，避免刷新过快占用资源
                    cpu = psutil.cpu_percent(interval=2)
                    self.cpu_signal.emit(cpu)
                else:
                    self.cpu_signal.emit(0.0)
                    time.sleep(5)
            except Exception:
                time.sleep(5)

# ================= 后台模型加载线程 (支持切换) =================
class ModelInitWorker(QThread):
    finished_signal = pyqtSignal(object, str) # 返回 实例, 模型名称

    def __init__(self, model_type="bge"):
        super().__init__()
        self.model_type = model_type # 'bge' or 'minilm'

    def run(self):
        print(f">>> 后台线程启动: 正在加载模型 [{self.model_type}]...")
        
        filter_instance = None
        name = ""
        
        if self.model_type == "bge":
            filter_instance = BGESemanticFilter()
            name = "BGE-Small-Zh-v1.5"
        else:
            filter_instance = LocalSemanticFilter()
            name = "MiniLM-L6 (Legacy)"
            
        self.finished_signal.emit(filter_instance, name)

# ================= 主界面 =================
class RAGRecallApp(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("RAG 工业级全流程 (Recall + RRF Fusion + Multi-Model Support)")
        self.resize(1400, 950) 
        self.setStyleSheet(STYLESHEET)
        
        self.settings = QSettings("MyCorp", "RAGRecall_Final_v12_BGE")
        self.cached_results = []
        self.cached_summary = ""
        self.cached_query = ""
        self.worker = None 
        
        # 本地小模型实例
        self.local_filter_instance = None
        self.current_model_name = "None"
        
        self.ensure_airline_dict()
        self.init_ui()
        
        # 启动系统监控
        self.sys_monitor = SystemMonitorWorker()
        self.sys_monitor.cpu_signal.connect(self.update_cpu_bar)
        self.sys_monitor.start()
        
        # 【修改点 1】: 默认自动启动加载 BGE (Win7 优化：延迟 1s 启动，防止 GUI 绘制竞争)
        self.log("🕒 正在准备后台加载默认模型 (BGE-Small)...")
        # 强制设置默认UI状态为 BGE
        self.radio_bge.setChecked(True)
        self.trigger_model_load("bge")
        
    def ensure_airline_dict(self):
        if not os.path.exists(AIRLINE_DICT_FILE):
            try:
                with open(AIRLINE_DICT_FILE, "w", encoding="utf-8") as f:
                    for airline in DEFAULT_AIRLINES:
                        f.write(airline + "\n")
            except:
                pass

    def init_ui(self):
        central_widget = QWidget()
        self.setCentralWidget(central_widget)
        
        main_layout = QHBoxLayout(central_widget)
        main_layout.setSpacing(10)
        
        # === Left Widget ===
        left_widget = QWidget()
        left_layout = QVBoxLayout(left_widget)
        
        # 1. DB & JSON Inputs
        db_layout = QHBoxLayout()
        self.db_path_edit = QLineEdit()
        self.db_path_edit.setText(self.settings.value("last_db_path", ""))
        btn_db = QPushButton("📂 向量库")
        btn_db.clicked.connect(self.browse_db)
        db_layout.addWidget(QLabel("Vector DB:"))
        db_layout.addWidget(self.db_path_edit)
        db_layout.addWidget(btn_db)
        left_layout.addLayout(db_layout)
        
        json_layout = QHBoxLayout()
        self.json_path_edit = QLineEdit()
        self.json_path_edit.setText(self.settings.value("last_json_path", ""))
        btn_json = QPushButton("📄 PageIndex")
        btn_json.clicked.connect(self.browse_json)
        json_layout.addWidget(QLabel("Structure JSON:"))
        json_layout.addWidget(self.json_path_edit)
        json_layout.addWidget(btn_json)
        left_layout.addLayout(json_layout)
        
        # 2. Search Mode & Model Selection
        mode_layout = QHBoxLayout()
        
        mode_layout.addWidget(QLabel("🔍 策略:"))
        self.mode_group = QButtonGroup(self)
        self.radio_smart = QRadioButton("🔵 智能")
        self.radio_smart.setObjectName("ModeSmart")
        self.radio_smart.setChecked(True)
        self.mode_group.addButton(self.radio_smart, 1)
        mode_layout.addWidget(self.radio_smart)
        
        self.radio_precise = QRadioButton("🟢 精准")
        self.radio_precise.setObjectName("ModePrecise")
        self.mode_group.addButton(self.radio_precise, 2)
        mode_layout.addWidget(self.radio_precise)
        
        self.radio_fuzzy = QRadioButton("🟡 模糊")
        self.radio_fuzzy.setObjectName("ModeFuzzy")
        self.mode_group.addButton(self.radio_fuzzy, 3)
        mode_layout.addWidget(self.radio_fuzzy)
        
        mode_layout.addSpacing(15)

        mode_layout.addWidget(QLabel("📂 文档偏好:"))
        self.combo_doc_type = QComboBox()
        self.combo_doc_type.addItems(["不指定类型", "数据表", "公司公文", "书籍/教材", "长篇论文", "技术文档", "法律条文", "LLM OCR文档", "LLM生成总结"])
        self.combo_doc_type.setFixedWidth(120)
        mode_layout.addWidget(self.combo_doc_type)

        mode_layout.addWidget(QLabel("🧠 LLM:"))
        self.combo_model = QComboBox()
        self.combo_model.addItems(["DeepSeek-R1", "DeepSeek-V3", "X1-70B-thinking", "X1-70B-fast"])
        self.combo_model.setFixedWidth(140)
        mode_layout.addWidget(self.combo_model)
        
        mode_layout.addStretch()
        left_layout.addLayout(mode_layout)

        # 3. Knowledge Config & Advanced Controls
        config_group = QGroupBox("🔧 知识库与本地模型配置 (Local Model)")
        config_layout = QVBoxLayout(config_group)
        
        # 3.1 Stopwords & Buttons
        sw_btn_layout = QHBoxLayout()
        self.stopwords_edit = QTextEdit()
        self.stopwords_edit.setPlaceholderText("在此输入停用词...")
        self.stopwords_edit.setMaximumHeight(35)
        sw_btn_layout.addWidget(self.stopwords_edit, 2)

        self.sw_name_edit = QLineEdit()
        self.sw_name_edit.setPlaceholderText("Config Name")
        self.sw_name_edit.setFixedWidth(80)
        sw_btn_layout.addWidget(self.sw_name_edit)
        
        btn_sw_save = QPushButton("💾 Save")
        btn_sw_save.clicked.connect(self.save_stopwords)
        btn_sw_save.setFixedHeight(28)
        sw_btn_layout.addWidget(btn_sw_save)
        
        btn_sw_load = QPushButton("📂 Load")
        btn_sw_load.clicked.connect(self.import_stopwords)
        btn_sw_load.setFixedHeight(28)
        sw_btn_layout.addWidget(btn_sw_load)
        
        btn_airline = QPushButton("✈️ 航司字典")
        btn_airline.clicked.connect(self.open_airline_editor)
        btn_airline.setFixedHeight(28)
        btn_airline.setStyleSheet("background-color: #8e44ad; color: white;")
        sw_btn_layout.addWidget(btn_airline)
        config_layout.addLayout(sw_btn_layout)

        # 3.2 Local Model Selection (New Feature)
        model_select_layout = QHBoxLayout()
        
        # 总开关
        self.chk_local_model = QCheckBox("🧠 启用本地小模型过滤")
        self.chk_local_model.setChecked(True)
        self.chk_local_model.setStyleSheet("color: #ffa502; font-weight: bold;")
        self.chk_local_model.setToolTip("使用本地 Embedding 模型进行粗排")
        self.chk_local_model.stateChanged.connect(self.toggle_local_model_ui)
        model_select_layout.addWidget(self.chk_local_model)
        
        # 分隔线
        line = QFrame()
        line.setFrameShape(QFrame.VLine)
        line.setFrameShadow(QFrame.Sunken)
        model_select_layout.addWidget(line)
        
        # 模型单选 (互斥)
        self.bg_models = QButtonGroup(self)
        
        self.radio_bge = QRadioButton("BGE-Small-Zh (New)")
        self.radio_bge.setToolTip("基于 D:\\Models\\bge-small-zh-v1.5，效果更好")
        # 默认选中 BGE
        self.radio_bge.setChecked(True)
        self.bg_models.addButton(self.radio_bge, 1)
        model_select_layout.addWidget(self.radio_bge)
        
        self.radio_minilm = QRadioButton("MiniLM-L6 (Legacy)")
        self.radio_minilm.setToolTip("旧版轻量级模型")
        self.bg_models.addButton(self.radio_minilm, 2)
        model_select_layout.addWidget(self.radio_minilm)
        
        # 连接切换信号
        self.bg_models.buttonClicked.connect(self.on_model_radio_clicked)

        model_select_layout.addStretch()
        
        # CPU 监控条 (Win7 保护)
        self.lbl_cpu = QLabel("CPU: 0%")
        model_select_layout.addWidget(self.lbl_cpu)
        self.progress_cpu = QProgressBar()
        self.progress_cpu.setRange(0, 100)
        self.progress_cpu.setFixedWidth(100)
        self.progress_cpu.setTextVisible(False)
        model_select_layout.addWidget(self.progress_cpu)

        config_layout.addLayout(model_select_layout)
        
        # 3.3 Advanced Parameters
        adv_layout = QHBoxLayout()
        self.chk_use_faiss = QCheckBox("⚡ 启用 FAISS 加速")
        self.chk_use_faiss.setChecked(True)
        self.chk_use_faiss.setStyleSheet("color: #00e676;") 
        adv_layout.addWidget(self.chk_use_faiss)
        
        # =======================================================
        # 【新增功能】文件路由建议开关
        # =======================================================
        self.chk_routing = QCheckBox("🔀 开启文件路由建议")
        self.chk_routing.setChecked(False) # 默认不开启
        self.chk_routing.setStyleSheet("color: #ff9ff3; font-weight: bold;")
        self.chk_routing.setToolTip("根据查询内容自动建议合适的数据源")
        adv_layout.addWidget(self.chk_routing)
        
        adv_layout.addStretch()
        
        adv_layout.addWidget(QLabel("向量召回数:"))
        self.spin_vector_k = QSpinBox()
        self.spin_vector_k.setRange(5, 50)
        self.spin_vector_k.setValue(10)
        self.spin_vector_k.setFixedWidth(50)
        adv_layout.addWidget(self.spin_vector_k)

        adv_layout.addWidget(QLabel("硬查询数:"))
        self.spin_json_k = QSpinBox()
        self.spin_json_k.setRange(5, 20)
        self.spin_json_k.setValue(10)
        self.spin_json_k.setFixedWidth(50)
        adv_layout.addWidget(self.spin_json_k)

        config_layout.addLayout(adv_layout)
        
        left_layout.addWidget(config_group)

        # 4. Query Input
        left_layout.addWidget(QLabel("用户查询 (Query):"))
        self.query_input = QTextEdit()
        self.query_input.setPlaceholderText("请输入问题...")
        self.query_input.setMaximumHeight(60)
        left_layout.addWidget(self.query_input)
        
        # 5. Search & Stop Buttons
        btn_layout = QHBoxLayout()
        self.btn_search = QPushButton("🚀 执行全流程 (Recall -> RRF -> Summary)")
        self.btn_search.setFixedHeight(45)
        self.btn_search.setStyleSheet("background-color: #2da44e; font-size: 15px;")
        self.btn_search.clicked.connect(self.start_recall)
        btn_layout.addWidget(self.btn_search)
        
        self.btn_stop = QPushButton("🛑 停止运行")
        self.btn_stop.setObjectName("StopBtn")
        self.btn_stop.setFixedHeight(45)
        self.btn_stop.setEnabled(False)
        self.btn_stop.clicked.connect(self.stop_recall)
        btn_layout.addWidget(self.btn_stop)
        left_layout.addLayout(btn_layout)
        
        # 6. Result Displays
        left_layout.addWidget(QLabel("🤖 智能总结 (Thinking + Answer):"))
        self.summary_display = QTextEdit()
        self.summary_display.setReadOnly(True)
        self.summary_display.setStyleSheet("""
            QTextEdit {
                background-color: #252526; 
                color: #dcdcaa; 
                font-family: 'Segoe UI', sans-serif; 
                font-size: 14px; 
                border: 1px solid #007acc;
                line-height: 1.6;
            }
        """)
        self.summary_display.setMinimumHeight(250)
        left_layout.addWidget(self.summary_display)

        left_layout.addWidget(QLabel("📚 RRF Fused Context:"))
        self.result_display = QTextEdit()
        self.result_display.setReadOnly(True)
        self.result_display.setStyleSheet("font-family: Consolas; font-size: 12px; color: #aaddff;")
        left_layout.addWidget(self.result_display)
        
        # 7. Export Area
        export_layout = QHBoxLayout()
        export_layout.addWidget(QLabel("导出格式:"))
        self.combo_format = QComboBox()
        self.combo_format.addItems(["xlsx", "csv", "txt", "docx", "md"])
        self.combo_format.setFixedWidth(100)
        export_layout.addWidget(self.combo_format)
        
        self.btn_export = QPushButton("💾 导出结果")
        self.btn_export.setStyleSheet("background-color: #d2691e;")
        self.btn_export.clicked.connect(self.export_data)
        export_layout.addWidget(self.btn_export)
        export_layout.addStretch() 
        left_layout.addLayout(export_layout)

        # === Right Console ===
        right_widget = QWidget()
        right_layout = QVBoxLayout(right_widget)
        right_layout.addWidget(QLabel("📟 System Console"))
        self.console_output = QTextEdit()
        self.console_output.setReadOnly(True)
        self.console_output.setStyleSheet("background-color: #111; color: #0f0; font-family: Consolas;")
        right_layout.addWidget(self.console_output)
        
        splitter = QSplitter(Qt.Horizontal)
        splitter.addWidget(left_widget)
        splitter.addWidget(right_widget)
        splitter.setSizes([900, 400])
        main_layout.addWidget(splitter)
        
        last_sw = self.settings.value("last_stopwords", "")
        if last_sw:
            self.stopwords_edit.setText(last_sw)

    # ================= 逻辑方法 =================
    
    def update_cpu_bar(self, usage):
        """更新 CPU 进度条，Win7 物理保护警示"""
        self.lbl_cpu.setText(f"CPU: {usage:.1f}%")
        self.progress_cpu.setValue(int(usage))
        
        # 颜色警示
        if usage < 50:
            style = "QProgressBar::chunk { background-color: #2da44e; }" # Green
        elif usage < 80:
            style = "QProgressBar::chunk { background-color: #f1c40f; }" # Yellow
        else:
            style = "QProgressBar::chunk { background-color: #e74c3c; }" # Red
            
        self.progress_cpu.setStyleSheet(style + "QProgressBar { border: 1px solid #555; border-radius: 3px; background-color: #333; }")

    def toggle_local_model_ui(self, state):
        enabled = (state == Qt.Checked)
        self.radio_bge.setEnabled(enabled)
        self.radio_minilm.setEnabled(enabled)
        
        if enabled:
            # 如果启用了，但当前没有加载模型，触发加载当前选中的
            if not self.local_filter_instance:
                selected_model = "bge" if self.radio_bge.isChecked() else "minilm"
                self.trigger_model_load(selected_model)
        else:
            self.log("⚪ 本地模型过滤已禁用。")

    def on_model_radio_clicked(self, btn):
        if btn == self.radio_bge:
            target = "bge"
        else:
            target = "minilm"
        
        # 只有当真正切换了类型才重载
        if target == "bge" and "BGE" in self.current_model_name:
            return
        if target == "minilm" and "MiniLM" in self.current_model_name:
            return

        self.trigger_model_load(target)

    def trigger_model_load(self, model_type):
        """触发后台线程加载，并处理互斥逻辑"""
        self.btn_search.setEnabled(False)
        self.log(f"🔄 正在切换至模型: {model_type.upper()} ...")
        
        # 1. 显式卸载旧模型，释放内存 (Win7 关键)
        if self.local_filter_instance:
            if hasattr(self.local_filter_instance, 'unload'):
                self.local_filter_instance.unload()
            del self.local_filter_instance
            self.local_filter_instance = None
            gc.collect() # 强制垃圾回收
            
        # 2. 启动加载线程
        self.model_loader = ModelInitWorker(model_type)
        self.model_loader.finished_signal.connect(self.on_model_loaded)
        self.model_loader.start()

    def on_model_loaded(self, instance, name):
        self.local_filter_instance = instance
        self.current_model_name = name
        self.btn_search.setEnabled(True)
        
        if instance and instance.is_loaded:
            self.log(f"✅ 模型就绪: {name}")
            # 自动选中对应的 radio (防止启动时的不一致)
            if "BGE" in name:
                self.radio_bge.setChecked(True)
            else:
                self.radio_minilm.setChecked(True)
        else:
            self.log(f"❌ 模型加载失败: {name}")
            self.chk_local_model.setChecked(False) # 自动关闭功能

    def browse_db(self):
        path, _ = QFileDialog.getOpenFileName(self, "选择数据库", "", "SQLite DB (*.db);;All Files (*.*)")
        if path:
            self.db_path_edit.setText(path)
            self.settings.setValue("last_db_path", path)

    def browse_json(self):
        path, _ = QFileDialog.getOpenFileName(self, "选择 JSON", "", "JSON Files (*.json);;All Files (*.*)")
        if path:
            self.json_path_edit.setText(path)
            self.settings.setValue("last_json_path", path)

    def log(self, msg):
        self.console_output.append(msg)
        cursor = self.console_output.textCursor()
        cursor.movePosition(QTextCursor.End)
        self.console_output.setTextCursor(cursor)

    def update_summary(self, text):
        self.cached_summary = text 
        self.summary_display.setMarkdown(text) 
        cursor = self.summary_display.textCursor()
        cursor.movePosition(QTextCursor.End)
        self.summary_display.setTextCursor(cursor)

    def display_results(self, results):
        self.cached_results = results 
        html = ""
        for item in results:
            score_text = f"{item['final_score']:.4f}"
            debug_info = item.get('debug_score', '')
            local_score = item.get('local_score', None)
            source = item.get('source', 'VECTOR')
            
            if "JSON" in source:
                source_style = "background-color: #00f260; color: #111; padding: 2px 5px; border-radius: 3px; font-weight: bold;"
            else:
                source_style = "background-color: #007acc; color: white; padding: 2px 5px; border-radius: 3px;"
            
            extra_badges = ""
            if local_score is not None:
                # 区分不同模型的显示颜色
                if "BGE" in self.current_model_name:
                    badge_color = "#e67e22" # Orange for BGE
                else:
                    badge_color = "#9b59b6" # Purple for MiniLM
                extra_badges += f" <span style='background-color:{badge_color};color:white;padding:1px 4px;border-radius:3px;'>Sim: {local_score:.3f}</span>"
            
            score_color = "#00ff00" 
            
            html += f"""
            <div style='border-bottom: 1px solid #555; padding: 12px; margin-bottom: 8px;'>
                <span style='color: #888; font-weight:bold;'>Rank #{item['rank']}</span> | 
                <span style='{source_style}'>{source}</span> | 
                <span style='color: {score_color}; font-weight: bold;'>RRF Score: {score_text}</span> 
                {extra_badges}
                <span style='color: #aaa; font-size:11px;'>[{debug_info}]</span><br>
                <div style='margin-top:5px; color: #ffcc00;'><b>[Section Path]</b> {item['path']}</div>
                <div style='margin-top:5px; background-color: #222; padding: 8px; border-left: 3px solid #2da44e; white-space: pre-wrap;'>
{item['content'][:200]}...
                </div>
            </div>
            """
        self.result_display.setHtml(html)

    def open_airline_editor(self):
        editor = AirlineDictEditor(self)
        editor.exec_()
    
    def get_current_stopwords(self):
        text = self.stopwords_edit.toPlainText().strip()
        if not text: return []
        return [w.strip() for w in text.replace('，', ',').split(',') if w.strip()]

    def get_airline_list(self):
        airlines = []
        if os.path.exists(AIRLINE_DICT_FILE):
            try:
                with open(AIRLINE_DICT_FILE, 'r', encoding='utf-8') as f:
                    airlines = [line.strip() for line in f if line.strip()]
            except: pass
        return airlines

    def save_stopwords(self):
        name = self.sw_name_edit.text().strip()
        if not name: return
        sw_list = self.get_current_stopwords()
        filename = f"stopwords_{name}.json"
        try:
            with open(filename, 'w', encoding='utf-8') as f:
                json.dump({"name": name, "stopwords": sw_list}, f, ensure_ascii=False, indent=2)
            self.log(f"💾 Stopwords profile '{name}' saved.")
        except Exception as e: self.log(f"❌ Save failed: {str(e)}")

    def import_stopwords(self):
        path, _ = QFileDialog.getOpenFileName(self, "Import Stopwords JSON", "", "JSON Files (*.json);;All Files (*.*)")
        if path:
            try:
                with open(path, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                sw_list = data if isinstance(data, list) else data.get("stopwords", [])
                self.stopwords_edit.setText(", ".join(sw_list))
                self.log(f"📂 Imported {len(sw_list)} words.")
            except Exception as e: self.log(f"❌ Import failed: {str(e)}")

    # ================= 核心流程控制 =================
    
    # 辅助方法：调用 API 获取 Query 向量
    def get_embedding_from_api(self, text):
        headers = {
            'Content-Type': 'application/json',
            'Authorization': f'Bearer {ROUTING_API_KEY}'
        }
        data = {
            "model": ROUTING_MODEL_NAME,
            "input": [text[:8000]] # 截断保护
        }
        try:
            self.log("📡 [Routing] 正在调用 BGE-M3 API 进行向量化...")
            response = requests.post(
                ROUTING_EMBED_URL, 
                headers=headers, 
                data=json.dumps(data), 
                verify=False,
                timeout=15
            )
            if response.status_code == 200:
                res = response.json()
                if "data" in res and len(res["data"]) > 0:
                    return res["data"][0]["embedding"]
            else:
                self.log(f"⚠️ API 错误: {response.status_code}")
        except Exception as e:
            self.log(f"⚠️ API 网络异常: {str(e)}")
        return None

    # 【Fixed】智能路由逻辑：修复硬匹配过宽及向量计算无归一化的问题
    def execute_routing_logic(self, query):
        # 确保路径指向 Step1 的总结库
        summary_db_path = os.path.join(os.getcwd(), "rooter_Step1", "doc_summary_v4.db")
        
        if not os.path.exists(summary_db_path):
            self.log(f"⚠️ 路由数据库不存在: {summary_db_path}")
            return
            
        suggested_file = None
        match_type = ""
        
        try:
            conn = sqlite3.connect(summary_db_path)
            cursor = conn.cursor()
            
            # --- Phase 1: 严格硬查询 (仅匹配文件名，避免内容干扰) ---
            # 修复：原逻辑同时匹配内容，导致"服务"等常用词总是命中简历文件
            sql_hard = "SELECT file_name FROM document_summaries WHERE file_name LIKE ? LIMIT 1"
            cursor.execute(sql_hard, (f"%{query}%",))
            res = cursor.fetchone()
            
            if res:
                suggested_file = res[0]
                match_type = "文件名硬匹配"
                self.log(f"🎯 发现文件名包含关键词，直接命中: {suggested_file}")
            else:
                # --- Phase 2: 向量语义匹配 (全量比对) ---
                query_vec = self.get_embedding_from_api(query)
                if query_vec:
                    cursor.execute("SELECT file_name, summary_vector FROM document_summaries")
                    rows = cursor.fetchall()
                    
                    best_score = -1.0
                    best_file = None
                    
                    # 关键修复：转为 float32 并计算范数
                    q_vec = np.array(query_vec, dtype=np.float32)
                    q_norm = np.linalg.norm(q_vec)
                    
                    self.log(f"📊 [Routing Debug] 开始计算语义相似度，库文件数: {len(rows)}")
                    
                    for fname, vec_str in rows:
                        if not vec_str or len(vec_str) < 10: continue
                        
                        try:
                            # 严格解析 JSON 向量
                            d_vec = np.array(json.loads(vec_str), dtype=np.float32)
                            d_norm = np.linalg.norm(d_vec)
                            
                            if d_norm == 0: continue
                            
                            # 关键修复：标准余弦相似度计算 (Dot / (Norm * Norm))
                            score = np.dot(q_vec, d_vec) / (q_norm * d_norm)
                            
                            # 输出每个文件的比对分，便于排查
                            self.log(f"   -> 检查文件: {fname[:30]}... | Score: {score:.4f}")
                            
                            if score > best_score:
                                best_score = score
                                best_file = fname
                        except Exception as e:
                            self.log(f"      解析向量失败({fname[:10]}): {e}")
                            continue
                    
                    # 设定阈值 (0.35) 避免噪音
                    if best_file and best_score > 0.35:
                        suggested_file = best_file
                        match_type = f"向量语义路由 (最高分: {best_score:.4f})"
                    else:
                        self.log(f"💡 [Routing] 最高相似度 {best_score:.4f} 低于阈值(0.35)，无建议。")
                else:
                    self.log("⚠️ 无法获取 Query 向量")

            conn.close()

            if suggested_file:
                self.log(f"💡 [Routing Result] 匹配到最相关文件: {suggested_file}")
                # 弹窗提示用户
                msg = f"🔍 路由系统分析建议：\n\n查询词: '{query}'\n匹配类型: {match_type}\n\n📂 建议切换至以下数据源进行查询：\n{suggested_file}"
                QMessageBox.information(self, "智能路由建议", msg)
            else:
                self.log("💡 [Routing] 未找到足够相关的历史文件总结。")
                
        except Exception as e:
            self.log(f"❌ [Routing Error] 执行异常: {str(e)}")

    def start_recall(self):
        db_path = self.db_path_edit.text().strip()
        json_path = self.json_path_edit.text().strip()
        query = self.query_input.toPlainText().strip()
        
        if not db_path or not os.path.exists(db_path):
            QMessageBox.warning(self, "Error", "无效的数据库路径")
            return
        if not query: return

        # 【新增功能】如果开启了路由，先执行路由逻辑
        if self.chk_routing.isChecked():
            self.execute_routing_logic(query)

        try:
            mode_id = self.mode_group.checkedId()
            search_mode = {1: "smart", 2: "precise", 3: "fuzzy"}.get(mode_id, "smart")

            summary_model = self.combo_model.currentText()
            doc_type = self.combo_doc_type.currentText()
            stopwords = self.get_current_stopwords()
            airline_names = self.get_airline_list()
            
            use_faiss = self.chk_use_faiss.isChecked()
            enable_local_filter = self.chk_local_model.isChecked()
            
            # 安全检查：如果勾选了但模型没加载成功，强制视为未启用
            if enable_local_filter and (not self.local_filter_instance or not self.local_filter_instance.is_loaded):
                self.log("⚠️ 警告: 本地模型选中但未加载成功，将跳过过滤步骤。")
                enable_local_filter = False

            limit_vector = self.spin_vector_k.value()
            limit_json = self.spin_json_k.value()
            
            self.cached_query = query
            self.cached_results = []
            self.cached_summary = ""
            
            self.btn_search.setEnabled(False)
            self.btn_stop.setEnabled(True) 
            self.result_display.clear()
            self.summary_display.clear() 
            self.console_output.clear()
            
            self.log(f"🚀 初始化任务... | 策略: {search_mode} | 模型: {summary_model}")
            
            if enable_local_filter:
                self.log(f"🧠 本地过滤: [启用] 使用 {self.current_model_name}")
            else:
                self.log(f"⚪ 本地过滤: [禁用]")
                
            # 实例化后端 Worker
            self.worker = RecallWorker(
                query_text=query, 
                db_path=db_path, 
                json_path=json_path, 
                search_mode=search_mode, 
                summary_model=summary_model, 
                doc_type=doc_type, 
                stopwords=stopwords, 
                airline_names=airline_names,
                use_faiss=use_faiss,
                enable_local_filter=enable_local_filter,
                local_filter_instance=self.local_filter_instance, # 传入当前活动的多态实例
                limit_vector=limit_vector,
                limit_json=limit_json
            )
            
            self.worker.log_signal.connect(self.log)
            self.worker.result_signal.connect(self.display_results)
            self.worker.summary_signal.connect(self.update_summary) 
            self.worker.finish_signal.connect(self.on_finished)
            self.worker.start()

        except Exception as e:
            QMessageBox.critical(self, "System Error", f"启动失败: {str(e)}")
            self.btn_search.setEnabled(True)
            self.btn_stop.setEnabled(False)

    def stop_recall(self):
        if self.worker and self.worker.isRunning():
            self.log("🛑 用户点击停止...")
            self.worker.stop()
            self.btn_stop.setEnabled(False) 

    def on_finished(self, success):
        self.btn_search.setEnabled(True)
        self.btn_stop.setEnabled(False)
        self.btn_stop.setText("🛑 停止运行")
        if success: self.log("✅ 全流程结束")
        else: self.log("❌ 流程中断或错误")

    # ================= 导出功能 (保持不变) =================
    def export_data(self):
        if not self.cached_summary and not self.cached_results:
            QMessageBox.warning(self, "提示", "没有可导出的结果。")
            return
        fmt = self.combo_format.currentText()
        timestamp = time.strftime("%Y%m%d_%H%M%S")
        filename = f"export_{timestamp}.{fmt}"
        save_path = os.path.join(os.getcwd(), filename)
        
        try:
            self.log(f"💾 正在导出为 {fmt} ...")
            if fmt == "xlsx":
                if pd is None: raise ImportError("缺少 pandas")
                data_rows = []
                for item in self.cached_results:
                    data_rows.append({
                        "Rank": item['rank'],
                        "Source": item.get('source', 'VECTOR'),
                        "RRF Score": item['final_score'],
                        "Local Score": item.get('local_score', ''),
                        "Section Path": item['path'],
                        "Content": item['content']
                    })
                df = pd.DataFrame(data_rows)
                with pd.ExcelWriter(save_path, engine='openpyxl') as writer:
                    pd.DataFrame([["DeepSeek Summary"], [self.cached_summary], [""]]).to_excel(writer, sheet_name='Report', index=False, header=False)
                    df.to_excel(writer, sheet_name='Report', index=False, startrow=4)
            
            elif fmt == "csv":
                with open(save_path, 'w', newline='', encoding='utf-8-sig') as f:
                    writer = csv.writer(f)
                    writer.writerow(["Summary", self.cached_summary])
                    writer.writerow([])
                    writer.writerow(["Rank", "Source", "Score", "LocalSim", "Path", "Content"])
                    for item in self.cached_results:
                        writer.writerow([item['rank'], item.get('source'), f"{item['final_score']:.4f}", item.get('local_score',''), item['path'], item['content']])
            
            elif fmt == "txt" or fmt == "md":
                with open(save_path, 'w', encoding='utf-8') as f:
                    f.write(f"Query: {self.cached_query}\n\nSummary:\n{self.cached_summary}\n\nResults:\n")
                    for item in self.cached_results:
                        f.write(f"#{item['rank']} [{item.get('source')}] Score:{item['final_score']:.4f} Local:{item.get('local_score')}\nPath: {item['path']}\n{item['content']}\n---\n")

            elif fmt == "docx":
                if docx is None: raise ImportError("缺少 python-docx")
                doc = docx.Document()
                doc.add_heading('RAG Report', 0)
                doc.add_paragraph(f"Query: {self.cached_query}")
                doc.add_heading('Summary', 1)
                doc.add_paragraph(self.cached_summary)
                doc.add_heading('Results', 1)
                for item in self.cached_results:
                    p = doc.add_paragraph()
                    run = p.add_run(f"#{item['rank']} [{item.get('source')}] Score:{item['final_score']:.4f}")
                    run.bold = True
                    doc.add_paragraph(f"Path: {item['path']}").italic = True
                    doc.add_paragraph(item['content'])
                    doc.add_paragraph("-" * 20)
                doc.save(save_path)

            self.log(f"✅ 导出成功: {save_path}")
            QMessageBox.information(self, "成功", f"已导出至: {save_path}")
        except Exception as e:
            self.log(f"❌ 导出错误: {str(e)}")
            QMessageBox.critical(self, "错误", str(e))

if __name__ == "__main__":
    app = QApplication(sys.argv)
    window = RAGRecallApp()
    window.show()
    sys.exit(app.exec_())